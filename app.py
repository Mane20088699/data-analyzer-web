#!/usr/bin/env python3
"""Data Analyzer — Web Edition  (Flask + Tailwind CSS, localhost)

Concept discovery is powered by KeyBERT (all-MiniLM-L6-v2). Everything else in
this file is a *local reasoning layer* built around KeyBERT and spaCy — no
external APIs, no agents: preprocessing, candidate re-ranking, synonym merging,
clustering, dependency / causal / contradiction / assumption analysis, an
enriched knowledge graph and an executive-insight roll-up. Fully offline and
deterministic.
"""

# ── SSL bypass ────────────────────────────────────────────────────────────────
import ssl as _ssl
_ssl._create_default_https_context = _ssl._create_unverified_context

import httpx as _httpx
_orig_client = _httpx.Client.__init__
def _patched_client(self, *a, **kw): kw.setdefault("verify", False); _orig_client(self, *a, **kw)
_httpx.Client.__init__ = _patched_client
_orig_async = _httpx.AsyncClient.__init__
def _patched_async(self, *a, **kw): kw.setdefault("verify", False); _orig_async(self, *a, **kw)
_httpx.AsyncClient.__init__ = _patched_async

import requests as _req, urllib3 as _u3
_u3.disable_warnings()
_orig_req = _req.Session.request
def _patched_req(self, *a, **kw): kw.setdefault("verify", False); return _orig_req(self, *a, **kw)
_req.Session.request = _patched_req

# ── Standard imports ──────────────────────────────────────────────────────────
import os, re, sys, json, uuid, random, warnings, tempfile, zipfile, io, unicodedata
from pathlib import Path
from collections import Counter, defaultdict
from typing import List, Dict, Tuple, Optional

try:
    import ftfy
except Exception:
    ftfy = None
warnings.filterwarnings("ignore")

import numpy as np
import pandas as pd
import networkx as nx
import nltk
import pdfplumber
import spacy
from docx import Document as DocxDocument
from keybert import KeyBERT
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.cluster import AgglomerativeClustering
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from flask import Flask, request, jsonify, send_file, render_template, Response

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# Determinism: pin the few seeds anything stochastic might read.
random.seed(0)
np.random.seed(0)

# ── NLTK bootstrap ────────────────────────────────────────────────────────────
for _r in ["punkt", "punkt_tab", "stopwords"]:
    nltk.download(_r, quiet=True)
from nltk.corpus import stopwords as _sw
from nltk.tokenize import sent_tokenize
STOPWORDS = set(_sw.words("english"))

# ── Load local models ─────────────────────────────────────────────────────────
print("Loading NLP models…")
NLP: Optional[spacy.Language] = None
NLP_MODEL = "none"
for _m in ["en_core_web_trf", "en_core_web_lg", "en_core_web_md", "en_core_web_sm"]:
    try:
        NLP = spacy.load(_m); NLP_MODEL = _m
        print(f"  [OK] spaCy: {_m}"); break
    except OSError:
        pass
if NLP is None:
    os.system(f'"{sys.executable}" -m spacy download en_core_web_lg')
    NLP = spacy.load("en_core_web_lg"); NLP_MODEL = "en_core_web_lg"
NLP.max_length = 2_000_000

KW_MODEL = KeyBERT("all-MiniLM-L6-v2")
print("  [OK] KeyBERT: all-MiniLM-L6-v2")

# ── MarkItDown: convert every upload to Markdown before analysis ────────────────
# https://github.com/microsoft/markitdown — one converter for PDF, Office, HTML,
# CSV/JSON/XML, images, etc. We feed its Markdown into the existing NLP pipeline
# so structure (headings, tables, lists) survives into concept/entity extraction.
try:
    from markitdown import MarkItDown
    MARKITDOWN = MarkItDown()
    print("  [OK] MarkItDown: documents → Markdown")
except Exception as _md_exc:
    MARKITDOWN = None
    print(f"  [!] MarkItDown unavailable ({_md_exc}) — falling back to built-in readers")

# ── GLiNER: zero-shot NER for labels spaCy's fixed schema can't supply ──────────
# Fills the gaps (job titles, products, events) that PERSON/ORG/GPE miss. Runs
# locally, deterministic in eval mode. Person/Org/Location stay with spaCy to
# avoid double-counting in the entity table and knowledge graph.
GLINER_LABELS = ["job title", "product", "event"]
GLINER_THRESHOLD = 0.5
try:
    from gliner import GLiNER
    GLINER = GLiNER.from_pretrained("urchade/gliner_small-v2.1")
    GLINER.eval()
    print("  [OK] GLiNER: urchade/gliner_small-v2.1 (zero-shot NER)")
except Exception as _gl_exc:
    GLINER = None
    print(f"  [!] GLiNER unavailable ({_gl_exc}) — skipping typed-entity extraction")
print("Ready.\n")

# ── Patterns ──────────────────────────────────────────────────────────────────
PROBLEM_WORDS = {
    "fail","fails","failed","failure","error","errors","bug","bugs",
    "defect","issue","issues","problem","problems","crash","crashes",
    "broken","corrupted","conflict","conflicts","incompatible","mismatch",
    "missing","undefined","exception","deadlock","leak","overflow",
    "risk","hazard","unsafe","vulnerable","vulnerability","violation",
    "breach","invalid","incorrect","wrong","timeout","blocked","frozen",
    "deprecated","obsolete","flaw","weakness",
}

# Subset of PROBLEM_WORDS that escalate a finding to Critical severity. "exploit"
# is intentionally absent: as a verb it is overwhelmingly benign ("exploit a
# niche/resource/opportunity"), and the security sense is carried reliably by the
# others ("breach", "malware", "unauthorized", "vulnerability").
SECURITY_WORDS = {
    "vulnerability","vulnerable","breach","unsafe","hazard","leak",
    "deadlock","overflow","unauthorized","malware","attack","compromise","violation",
}

# Relationship phrases that express a dependency (for the dependency map).
DEP_REL_WORDS = {
    "depends on","depend","requires","require","needs","need","uses","use",
    "based on","relies on","rely","part of","composed of","contains","contain",
    "include","includes","built on","powered by","derived from",
}

# Opposing predicate pairs → flag contradictory relationships on the same pair.
ANTONYM_PAIRS = [
    ("increase","decrease"),("increases","decreases"),("enable","prevent"),
    ("enables","prevents"),("allow","block"),("allows","blocks"),
    ("cause","prevent"),("causes","prevents"),("support","contradict"),
    ("supports","contradicts"),("improve","reduce"),("improves","reduces"),
    ("add","remove"),("adds","removes"),("open","close"),("opens","closes"),
    ("accept","reject"),("accepts","rejects"),("start","stop"),("rise","fall"),
]

# Assumption markers → confidence that the sentence states a (possibly hidden) premise.
ASSUMPTION_MARKERS = {
    "it is assumed": 85, "we assume": 85, "taken for granted": 82, "assumption": 80,
    "assuming": 80, "presume": 78, "assume": 78, "presumably": 70, "supposedly": 68,
    "by default": 65, "given that": 60, "we believe": 60, "suppose": 60,
    "in theory": 58, "expected to": 55, "should be": 50,
}

_NEGATIONS = {
    "not","no","never","cannot","can't","won't","doesn't","don't","isn't",
    "aren't","wasn't","weren't","without","unable","none","neither","nor",
}

# Per-pattern confidence for causal/relational regexes (explicit causality > weak association).
_CAUSAL_CONF = {
    "causes": 88, "caused by": 86, "leads to": 85, "results in": 85, "due to": 84,
    "triggers": 82, "depends on": 80, "requires": 80, "prevents": 80, "enables": 78,
    "blocks": 78, "if → then": 75, "impacts": 70, "affects": 70, "demands": 66,
    "extracts": 60, "associated with": 60, "reveals": 58, "shows": 56, "shares": 55,
    "reflects": 55, "faces": 55, "recognizes": 55, "represents": 54, "carries": 52,
    "symbolizes": 50, "explores": 50, "enters": 50, "reclaims": 50, "escapes": 50,
}

CAUSAL_PATTERNS: List[Tuple[str, str]] = [
    (r"(\w[\w\s]{1,40}?)\s+causes?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "causes"),
    (r"(\w[\w\s]{1,40}?)\s+leads?\s+to\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "leads to"),
    (r"(\w[\w\s]{1,40}?)\s+results?\s+in\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "results in"),
    (r"(\w[\w\s]{1,40}?)\s+triggers?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "triggers"),
    (r"(\w[\w\s]{1,40}?)\s+depends?\s+on\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "depends on"),
    (r"(\w[\w\s]{1,40}?)\s+requires?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "requires"),
    (r"(\w[\w\s]{1,40}?)\s+affects?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "affects"),
    (r"(\w[\w\s]{1,40}?)\s+impacts?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "impacts"),
    (r"(\w[\w\s]{1,40}?)\s+prevents?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "prevents"),
    (r"(\w[\w\s]{1,40}?)\s+enables?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "enables"),
    (r"(\w[\w\s]{1,40}?)\s+blocks?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "blocks"),
    (r"if\s+(\w[\w\s]{1,40}?),?\s+then\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "if → then"),
    (r"(\w[\w\s]{1,40}?)\s+is\s+caused\s+by\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "caused by"),
    (r"(\w[\w\s]{1,40}?)\s+is\s+due\s+to\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "due to"),
    (r"(\w[\w\s]{1,40}?)\s+is\s+associated\s+with\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "associated with"),
    (r"(\w[\w\s]{1,40}?)\s+represents?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "represents"),
    (r"(\w[\w\s]{1,40}?)\s+symbolizes?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "symbolizes"),
    (r"(\w[\w\s]{1,40}?)\s+explores?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "explores"),
    (r"(\w[\w\s]{1,40}?)\s+faces?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "faces"),
    (r"(\w[\w\s]{1,40}?)\s+carries?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "carries"),
    (r"(\w[\w\s]{1,40}?)\s+reflects?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "reflects"),
    (r"(\w[\w\s]{1,40}?)\s+reveals?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "reveals"),
    (r"(\w[\w\s]{1,40}?)\s+shows?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "shows"),
    (r"(\w[\w\s]{1,40}?)\s+demands?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "demands"),
    (r"(\w[\w\s]{1,40}?)\s+extracts?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "extracts"),
    (r"(\w[\w\s]{1,40}?)\s+share[sd]?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "shares"),
    (r"(\w[\w\s]{1,40}?)\s+enters?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "enters"),
    (r"(\w[\w\s]{1,40}?)\s+reclaims?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "reclaims"),
    (r"(\w[\w\s]{1,40}?)\s+recognizes?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "recognizes"),
    (r"(\w[\w\s]{1,40}?)\s+escapes?\s+(\w[\w\s]{1,40}?)(?=[,.\n]|$)", "escapes"),
]

VARIABLE_PATTERNS: List[Tuple[str, str]] = [
    (r"\b([A-Z][a-zA-Z0-9_]{1,20})\s*=\s*([\d.]+\s*[a-zA-Z%°/]*)", "assignment"),
    (r"\b([a-zA-Z_]\w{1,20})\s+is\s+set\s+to\s+([\d.]+\s*[a-zA-Z%°/]*)", "configured"),
    (r"\b([a-zA-Z_]\w{1,20})\s*:\s*([\d.]+\s*[a-zA-Z%°/]*)\b", "defined"),
    (r"\b(threshold|limit|maximum|minimum|rate|value|count|size|capacity"
     r"|timeout|interval|budget|score|ratio|percentage|speed|temperature"
     r"|pressure|voltage|frequency|duration|weight|height|width|length)"
     r"\s+(?:of\s+[\w\s]{1,20}\s+)?(?:is\s+)?([\d.]+\s*(?:[a-zA-Z%°/]+)?)", "parameter"),
    (r"\b([\d.]+)\s*(percent|%|degrees?|ms|milliseconds?|seconds?|minutes?"
     r"|hours?|days?|meters?|km|kg|lb|MB|GB|TB|KB|RPM|fps|Hz|kHz|MHz|V|A|W)\b", "measurement"),
]

_SVO_DEPS = {"ROOT", "relcl", "xcomp", "advcl", "conj", "acl", "ccomp", "pcomp"}

_PERSON_PREFIXES = re.compile(
    r"^(O'|O’|Mc|Mac|De|Di|Van|Von|Le|La|St\.?\s)\w", re.IGNORECASE)

# spaCy NER labels that name proper nouns. In normal prose these are capitalised,
# so a fully lower-case span (e.g. "tae kwon") tagged as one is almost always a
# model misfire — we drop those rather than trust the statistical guess.
_PROPER_LABELS = {"PERSON", "ORG", "GPE", "LOC", "NORP", "FAC", "PRODUCT",
                  "EVENT", "WORK_OF_ART", "LANGUAGE"}

# spaCy numeric/value labels that aren't useful *named* entities ("four", "first",
# "23%", "$1.2M"). They're noise in the Entities table and are already captured
# with higher precision by extract_structured() → the Structured Data tab.
_NOISE_ENTITY_LABELS = {"CARDINAL", "ORDINAL", "PERCENT", "MONEY", "QUANTITY"}

# Citation-style DATE spans that swamp the entity table on academic PDFs: bare
# years ("1998"), year+suffix ("1991a"), year lists ("1998, 2003"), and journal
# vol:page locators ("3:613-31", "44:487"). Descriptive dates ("July 26, 2004",
# "100 years old", "three-year") don't match and are kept.
_NOISE_DATE_RE = re.compile(
    r"^(?:(?:19|20)\d{2}[a-z]?(?:[,;/&\s]+(?:19|20)?\d{1,4}[a-z]?)*"
    r"|\d{1,4}\s*:\s*\d{1,4}.*"
    r"|[\d.,;:/&\s\-–]+)$")

# Journal / publisher / citation abbreviations that NER mistakes for PERSON/ORG/GPE
# ("Ecol.", "Evol.", "Univ.", "Annu."). Most arrive via the bibliography; this is
# defence-in-depth for any that leak into the body via running heads.
_ENTITY_STOPWORDS = {
    "ecol", "evol", "biol", "syst", "annu", "rev", "univ", "herpetol", "zool",
    "sci", "nat", "conserv", "proc", "soc", "inst", "dep", "div", "natl", "acad",
    "ed", "eds", "vol", "pp", "fig", "no", "et al", "limnol", "oceanogr", "midl",
    "freshw", "copeia", "oikos", "ecosys", "geogr", "behav", "entomol", "pac",
    "gtr", "cv", "phd", "doi", "isbn",
}

def _is_noise_date(t: str) -> bool:
    return bool(_NOISE_DATE_RE.match(t.strip()))

def _is_stopword_entity(t: str) -> bool:
    low = re.sub(r"[.\s]+$", "", t.strip().lower())
    return low in _ENTITY_STOPWORDS

def _looks_proper(text: str) -> bool:
    """True if any letter is upper-case — real proper nouns are capitalised
    somewhere ("John", "iPhone", "eBay"); pure lower-case spans are misfires."""
    return any(c.isupper() for c in text)

# Pronouns and relativisers must never anchor an SVO triple. As a subject/object
# they collapse many unrelated sentences onto one meaningless hub ("they", "it",
# "that", "which") that then dominates the knowledge graph as the biggest node.
# spaCy tags most as PRON, but relative "that"/"which"/"who" sometimes come back
# as DET/other — so we also match by lemma.
_PRONOUN_LEMMAS = {
    "i", "you", "he", "she", "it", "we", "they", "me", "him", "her", "us", "them",
    "myself", "yourself", "himself", "herself", "itself", "ourselves",
    "yourselves", "themselves", "oneself",
    "this", "that", "these", "those", "which", "who", "whom", "whose", "what",
    "someone", "somebody", "something", "anyone", "anybody", "anything",
    "everyone", "everybody", "everything", "one", "ones", "other", "others",
    "none", "nobody", "nothing", "each", "either", "neither",
}

def _is_pronoun(tok) -> bool:
    """A subject/object token that should never become a graph node — a pronoun
    or relativiser carrying no entity identity of its own."""
    return tok.pos_ == "PRON" or tok.lemma_.lower() in _PRONOUN_LEMMAS

def _word_set(text: str) -> set:
    """Lower-case alphabetic word tokens of a string, for whole-word keyword
    matching (so "issue" never matches inside "tissue")."""
    return set(re.findall(r"[a-z]+", text.lower()))

_TRIVIAL_CONCEPTS = {
    "one","two","three","four","five","six","seven","eight","nine","ten",
    "first","second","third","fourth","fifth","last","next","new","old",
    "good","great","well","also","even","just","like","much","many","every",
    "said","say","told","tell","got","get","know","think","want","need",
    "go","come","see","look","take","give","make","use","find","turn",
    "day","time","man","woman","people","way","thing","place","part","year",
    "yes","no","not","yet","still","back","away","down","up","out","over",
    "never","always","already","once","again","around","together",
}

# ── Document readers ──────────────────────────────────────────────────────────
def read_pdf(path: str) -> str:
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: parts.append(t)
    raw = "\n\n".join(parts)
    raw = re.sub(r'(?<!\n)\n(?!\n)', ' ', raw)
    raw = re.sub(r' {2,}', ' ', raw)
    return raw.strip()

def read_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if r: parts.append(r)
    return "\n".join(parts)

def read_text(path: str) -> str:
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try: return Path(path).read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError): pass
    return Path(path).read_bytes().decode("ascii", errors="replace")

# Friendly format labels by extension (for the Files table).
_FORMAT_LABELS = {
    ".pdf": "PDF", ".docx": "Word Document", ".doc": "Word Document",
    ".pptx": "PowerPoint", ".ppt": "PowerPoint",
    ".xlsx": "Excel", ".xls": "Excel", ".csv": "CSV",
    ".html": "HTML", ".htm": "HTML", ".json": "JSON", ".xml": "XML",
    ".epub": "EPUB", ".md": "Markdown", ".markdown": "Markdown",
}

def read_document(path: str) -> Tuple[str, str]:
    ext = Path(path).suffix.lower()
    label = _FORMAT_LABELS.get(ext, "Text / ASCII")

    # Primary path: let MarkItDown convert the document to Markdown, then analyse
    # that. Markdown preserves headings/tables/lists the legacy readers flatten.
    if MARKITDOWN is not None:
        try:
            md = MARKITDOWN.convert(path).text_content
            if md and md.strip():
                return md, f"{label} → Markdown"
        except Exception:
            pass  # fall through to the built-in extractors below

    # Fallback: the original format-specific readers.
    if ext == ".pdf": return read_pdf(path), "PDF"
    if ext in (".docx", ".doc"): return read_docx(path), "Word Document"
    return read_text(path), label

# ── NLP helpers ───────────────────────────────────────────────────────────────
def chunk_text(text: str, max_chunk: int = 80_000) -> List[str]:
    if len(text) <= max_chunk: return [text]
    sents = sent_tokenize(text)
    chunks, cur = [], ""
    for s in sents:
        if len(cur) + len(s) > max_chunk and cur:
            chunks.append(cur.strip()); cur = s
        else:
            cur += " " + s
    if cur.strip(): chunks.append(cur.strip())
    return chunks or [text[:max_chunk]]

def _clean(text: str, n: int = 200) -> str:
    return re.sub(r"\s+", " ", text).strip()[:n]

_LEADING_FLUFF = re.compile(
    r"^(?:the|a|an|and|but|or|nor|so|then|thus|hence|therefore|however|this|that|"
    r"these|those|its|their|our|your|his|her|which|who|it|they)\s+", re.IGNORECASE)

def _node_label(text: str, n: int = 40) -> str:
    """Canonical graph-node label: trimmed + leading articles/conjunctions stripped.

    Collapses fragments like 'and the Database' / 'The Database' into 'Database'
    so causal chains, dependencies and problems don't fragment on surface form."""
    s = _clean(text, n)
    prev = None
    while prev != s:
        prev = s
        s = _LEADING_FLUFF.sub("", s).strip()
    return s

def _embed(texts: List[str]) -> np.ndarray:
    """Embed phrases with KeyBERT's own sentence-transformer — no extra model."""
    if not texts:
        return np.zeros((0, 384), dtype="float32")
    return np.asarray(KW_MODEL.model.embed(texts))

def _cluster(emb: np.ndarray, threshold: float = 0.55) -> List[int]:
    """Deterministic agglomerative clustering by cosine distance."""
    n = len(emb)
    if n <= 2:
        return [0] * n
    try:
        model = AgglomerativeClustering(
            n_clusters=None, metric="cosine", linkage="average",
            distance_threshold=threshold)
        return list(model.fit_predict(emb))
    except Exception:
        return [0] * n

_HEADING_RE = re.compile(r"^(?:\d+(?:\.\d+)*[.)]?\s+)?[A-Z][^.!?]{2,68}$")

def _headings(text: str) -> List[str]:
    """Heuristic heading lines (short, Title/UPPER case, no terminal period)."""
    out = []
    for line in text.splitlines():
        l = line.strip()
        if 3 <= len(l) <= 70 and len(l.split()) <= 10 and \
           (l.isupper() or _HEADING_RE.match(l)):
            out.append(l.lower())
    return out

def _tail_windows(text: str, head: int = 60_000, size: int = 20_000,
                  max_windows: int = 4) -> List[str]:
    """Sliding windows over the part of long docs KeyBERT's head pass misses."""
    if len(text) <= head:
        return []
    seg = text[head:head + size * max_windows]
    return [seg[i:i + size] for i in range(0, len(seg), size)]

def normalize_unicode(text: str) -> str:
    """Repair encoding damage and decompose typographic ligatures.
    ftfy fixes mojibake (CÃ´tÃ© → Côté); NFKC maps ligatures (ﬁ→fi, ﬂ→fl)
    that ftfy's NFC pass leaves intact."""
    if ftfy is not None:
        text = ftfy.fix_text(text)
    return unicodedata.normalize("NFKC", text)

# Reference/bibliography section headings. Specific phrases match case-insensitively
# (rare in prose); bare "REFERENCES" only in its uppercase heading form so a
# sentence like "see references above" doesn't trip it. PDF extraction often
# flattens the heading inline, so this is intentionally not line-anchored.
_REFERENCES_RE = re.compile(
    r"(?i:literature\s+cited|works\s+cited|bibliography|references\s+cited)"
    r"|\bREFERENCES\b")

def strip_references(text: str) -> str:
    """Drop a trailing references/bibliography section so cited-author names and
    journal abbreviations ('Ecol', 'Welsh HH') don't swamp entity/relationship
    extraction. Only cuts at a heading in the back half of the document, so a doc
    that merely mentions 'references' up front is left intact."""
    for m in _REFERENCES_RE.finditer(text):
        if m.start() >= len(text) * 0.5:
            return text[:m.start()].rstrip()
    return text

# Recurring PDF page furniture (download stamps, running heads, timestamps). These
# survive line-based de-boilerplating because markitdown flattens them inline, so
# we scrub them with substitutions on the whole text.
_PAGE_FURNITURE = [
    # Whole download stamp through "... on MM/DD/YY". One lazy span so it works
    # even when markitdown glues tokens ("...annualreviews.orgby U.S. Department
    # of Agriculture on 09/27/06") — a separate "by ..." rule would miss those.
    re.compile(r"Downloaded from\b[^\n]{0,140}?on\s+\d{1,2}/\d{1,2}/\d{2,4}\.?", re.I),
    # Journal running head / volume locator ("Annu. Rev. Ecol. ... 2004.35:405-434").
    re.compile(r"\b[A-Z][a-z]{2,5}\.\s+Rev\.[^\n]{0,60}?\d{4}\.\d+:\d+[-–]\d+\.?"),
    re.compile(r"\b\d{4}\.\d+:\d+[-–]\d+\b"),
    # "For personal use only." plus any page number that trails the footer, so the
    # next sentence doesn't inherit a stray "406"/"414" as its subject.
    re.compile(r"\bFor personal use only\.?\s*\d{0,4}", re.I),
    # Stray "DD Mon YYYY HH:MM <pagenum>" timestamps when present.
    re.compile(r"\b\d{1,2}\s+[A-Z][a-z]{2}\s+\d{4}\s+\d{1,2}:\d{2}(?:\s+\d{1,4})?\b"),
]

def strip_page_furniture(text: str) -> str:
    for rx in _PAGE_FURNITURE:
        text = rx.sub(" ", text)
    return text

_MD_TABLE_SEP = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?$")

def flatten_md_tables(text: str) -> str:
    """Turn markdown/extracted table rows into short sentences. Otherwise a whole
    pipe-delimited table flows into one sentence and becomes a giant SVO subject
    (e.g. 'Kingdom Animalia Phylum Chordata … evolve early tetrapods')."""
    out = []
    for line in text.split("\n"):
        s = line.strip()
        if _MD_TABLE_SEP.match(s):
            continue                                    # divider row → drop
        if s.startswith("|") and s.count("|") >= 2:
            cells = [c.strip() for c in s.strip("|").split("|") if c.strip()]
            out.append("; ".join(cells) + "." if cells else "")
        else:
            out.append(line)
    return "\n".join(out)

def preprocess_text(text: str) -> str:
    """Normalise text before it reaches KeyBERT/spaCy (the 'preprocessing' layer)."""
    text = normalize_unicode(text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"-\n(?=[a-z])", "", text)                 # de-hyphenate wraps
    text = strip_page_furniture(text)
    text = flatten_md_tables(text)
    lines = [l.rstrip() for l in text.split("\n")]
    # Drop repeated page headers/footers (same short line appearing many times).
    freq = Counter(l.strip() for l in lines if 0 < len(l.strip()) <= 60)
    boiler = {l for l, c in freq.items()
              if c >= 5 and not l.endswith((".", ":", "?", "!"))}
    if boiler:
        lines = [l for l in lines if l.strip() not in boiler]
    text = "\n".join(lines)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

# ── Core spaCy pass: sentences + entities + SVO in one sweep ────────────────────
def parse_document(text: str):
    """One spaCy pass → (sentences, entity counter, entity→sentence map, SVO triples)."""
    sents: List[str] = []
    ent_counter: Counter = Counter()
    ent_sents: Dict[str, set] = defaultdict(set)
    svo: List[Dict] = []

    for chunk in chunk_text(text, 40_000):
        doc = NLP(chunk)
        for sent in doc.sents:
            stext = _clean(sent.text, 100_000)
            if not stext:
                continue
            si = len(sents)
            sents.append(stext)

            for ent in sent.ents:
                t = ent.text.strip(); label = ent.label_
                if len(t) <= 1: continue
                # Drop bare numerals/values (CARDINAL "four", PERCENT, MONEY…) —
                # noise here, and already in the Structured Data tab.
                if label in _NOISE_ENTITY_LABELS:
                    continue
                # Drop citation-year DATE noise ("1998", "1991a", "3:613-31").
                if label == "DATE" and _is_noise_date(t):
                    continue
                # Drop journal/citation abbreviations NER reads as names ("Ecol.").
                if _is_stopword_entity(t):
                    continue
                # Drop lower-case proper-noun misfires (e.g. "tae kwon" → PERSON).
                if label in _PROPER_LABELS and not _looks_proper(t):
                    continue
                if label == "GPE" and (_PERSON_PREFIXES.match(t) or "'" in t):
                    label = "PERSON"
                ent_counter[(t, label)] += 1
                ent_sents[t.lower()].add(si)

            for tok in sent:
                if tok.pos_ != "VERB" or tok.dep_ not in _SVO_DEPS: continue
                verb = tok.lemma_.lower()
                subjs = [w for w in tok.lefts  if w.dep_ in ("nsubj","nsubjpass","csubj")
                         and not _is_pronoun(w)]
                objs  = [w for w in tok.rights if w.dep_ in ("dobj","attr","acomp")
                         and not _is_pronoun(w)]
                for prep in tok.rights:
                    if prep.dep_ == "prep":
                        objs += [c for c in prep.children
                                 if c.dep_ == "pobj" and not _is_pronoun(c)]
                for s in subjs:
                    st = _clean(" ".join(w.text for w in s.subtree
                                         if w.dep_ not in ("det","punct") and len(w.text) > 1))
                    for o in objs:
                        ot = _clean(" ".join(w.text for w in o.subtree
                                             if w.dep_ not in ("det","punct") and len(w.text) > 1))
                        if st and ot and len(st) > 2:
                            svo.append({"Subject": st, "Relationship": verb, "Object": ot,
                                        "Rel. Type": "syntactic (SVO)", "Confidence": 65,
                                        "Source Sentence": _clean(sent.text, 250), "_si": si})

    seen, dedup = set(), []
    for r in svo:
        key = (r["Subject"].lower()[:40], r["Relationship"][:20], r["Object"].lower()[:40])
        if key not in seen:
            seen.add(key); dedup.append(r)
    return sents, ent_counter, ent_sents, dedup[:300]

# ── Concept extraction: KeyBERT engine + reasoning layer ────────────────────────
def extract_concepts(text: str, sents: List[str], ent_sents: Dict[str, set]) -> List[Dict]:
    """KeyBERT discovers candidates; the surrounding layer filters, merges,
    re-ranks, clusters, describes and links them. KeyBERT stays the engine."""
    try:
        # 1 ── candidate discovery (KeyBERT): head pass + sliding tail windows
        cand: Dict[str, float] = {}
        def _harvest(segment: str, weight: float, n: int):
            segment = segment.strip()
            if len(segment) < 40:
                return
            for kw, sc in KW_MODEL.extract_keywords(
                    segment, keyphrase_ngram_range=(1, 3), stop_words="english",
                    top_n=n, use_mmr=True, diversity=0.5):
                cand[kw] = max(cand.get(kw, 0.0), sc * weight)

        _harvest(text[:60_000], 1.0, 40)
        for w in _tail_windows(text):
            _harvest(w, 0.85, 10)

        # 2 ── filter trivial / numeric / too-short candidates
        cands: List[Tuple[str, float]] = []
        for kw, sc in cand.items():
            words = kw.lower().split()
            if len(words) == 1 and words[0] in _TRIVIAL_CONCEPTS: continue
            if re.fullmatch(r'[\d\s,.\-]+', kw): continue
            if len(words) == 1 and len(kw) <= 2: continue
            if all(w in STOPWORDS for w in words): continue
            cands.append((kw, sc))
        if not cands:
            return []

        phrases = [k for k, _ in cands]
        emb = _embed(phrases)

        # 3 ── merge near-duplicate / synonymous phrases (cosine ≥ 0.82 or substring)
        sim = cosine_similarity(emb) if len(phrases) > 1 else np.array([[1.0]])
        order = sorted(range(len(phrases)), key=lambda i: -cands[i][1])
        merged_into: Dict[int, int] = {}
        for pos, i in enumerate(order):
            if i in merged_into: continue
            for j in order[pos + 1:]:
                if j in merged_into: continue
                a, b = phrases[i].lower(), phrases[j].lower()
                if sim[i][j] >= 0.82 or a == b or a in b or b in a:
                    merged_into[j] = i
        keep = [i for i in range(len(phrases)) if i not in merged_into]

        # 4 ── multi-signal feature scoring → 0..100 importance
        low = text.lower()
        heads = " || ".join(_headings(text))
        sents_low = [s.lower() for s in sents]
        kb_scores = [cands[i][1] for i in keep]
        kb_min, kb_max = min(kb_scores), max(kb_scores)

        rows: List[Dict] = []
        for i in keep:
            kw = phrases[i]; klow = kw.lower()
            kb = cands[i][1]
            kb_n = (kb - kb_min) / (kb_max - kb_min) if kb_max > kb_min else 1.0
            freq = low.count(klow)
            freq_n = min(np.log1p(freq) / np.log1p(20), 1.0)
            pos = low.find(klow)
            pos_n = 1.0 - (pos / len(low)) if pos >= 0 and low else 0.0
            head_n = 1.0 if klow and klow in heads else 0.0
            csents = {si for si, s in enumerate(sents_low) if klow in s}
            ent_hits = [(ename, len(csents & eset))
                        for ename, eset in ent_sents.items()
                        if len(ename) > 2 and (csents & eset)]
            ent_n = min(len(ent_hits) / 5.0, 1.0)
            importance = 100 * (0.42*kb_n + 0.24*freq_n + 0.14*pos_n +
                                0.10*head_n + 0.10*ent_n)
            rel_ents = [e for e, _ in sorted(ent_hits, key=lambda x: -x[1])[:3]]
            desc = ""
            for si in sorted(csents):
                if 30 <= len(sents[si]) <= 240:
                    desc = sents[si]; break
            if not desc and csents:
                desc = sents[min(csents)]
            rows.append({"_i": i, "Concept": kw, "kb": round(float(kb), 4),
                         "imp": round(float(importance), 1), "freq": int(freq),
                         "ents": ", ".join(rel_ents), "desc": _clean(desc, 200)})

        # 5 ── cluster concepts by embedding → name each theme after its top concept
        kept_emb = emb[[r["_i"] for r in rows]]
        labels = _cluster(kept_emb)
        best_in: Dict[int, Tuple[float, str]] = {}
        for r, lab in zip(rows, labels):
            if r["imp"] > best_in.get(lab, (-1.0, ""))[0]:
                best_in[lab] = (r["imp"], r["Concept"])

        # 6 ── rank, assign Core/Supporting/Derived tier, emit ordered records
        ranked = sorted(zip(rows, labels), key=lambda rl: -rl[0]["imp"])
        n = len(ranked)
        out: List[Dict] = []
        for rank, (r, lab) in enumerate(ranked):
            tier = ("Core" if rank < max(1, n * 0.3)
                    else "Supporting" if rank < max(2, n * 0.7) else "Derived")
            out.append({
                "Concept": r["Concept"], "Type": tier, "Importance": r["imp"],
                "Relevance Score": r["kb"], "Theme / Cluster": best_in[lab][1],
                "Frequency": r["freq"], "Related Entities": r["ents"],
                "Description": r["desc"],
            })
        return out[:40]
    except Exception as exc:
        return [{"Concept": f"Error: {exc}", "Type": "", "Importance": 0,
                 "Relevance Score": 0.0, "Theme / Cluster": "", "Frequency": 0,
                 "Related Entities": "", "Description": ""}]

# ── Entity extraction: spaCy NER + alias merge + importance/tiers ───────────────
def build_entities(ent_counter: Counter, ent_sents: Dict[str, set],
                   sents: List[str], concepts: List[Dict]) -> List[Dict]:
    raw = [{"text": t, "type": l, "count": c} for (t, l), c in ent_counter.items()]
    raw.sort(key=lambda e: (-len(e["text"]), -e["count"]))

    # Merge aliases of the same type when one name's tokens subset another's
    # (e.g. "Smith" → "John Smith", "Apple" → "Apple Inc").
    canon: List[Dict] = []
    for e in raw:
        toks = set(e["text"].lower().split())
        hit = None
        for c in canon:
            if c["type"] != e["type"]: continue
            ctoks = set(c["text"].lower().split())
            if toks and (toks <= ctoks or ctoks <= toks):
                hit = c; break
        if hit:
            hit["count"] += e["count"]
            hit["sent"] |= ent_sents.get(e["text"].lower(), set())
        else:
            canon.append({"text": e["text"], "type": e["type"], "count": e["count"],
                          "sent": set(ent_sents.get(e["text"].lower(), set()))})

    concept_terms = [c["Concept"].lower() for c in concepts[:20] if c.get("Concept")]
    sents_low = [s.lower() for s in sents]
    maxc = max((c["count"] for c in canon), default=1)

    scored: List[Dict] = []
    for c in canon:
        freq_n = c["count"] / maxc
        overlap = sum(1 for si in c["sent"] if si < len(sents_low)
                      for t in concept_terms if t and t in sents_low[si])
        conc_n = min(overlap / 6.0, 1.0)
        importance = round(100 * (0.6 * freq_n + 0.4 * conc_n), 1)
        desc = ""
        for si in sorted(c["sent"]):
            if si < len(sents) and 25 <= len(sents[si]) <= 240:
                desc = sents[si]; break
        scored.append({"text": c["text"], "type": c["type"], "count": c["count"],
                       "imp": importance, "desc": _clean(desc, 200)})

    scored.sort(key=lambda e: (-e["imp"], -e["count"]))
    n = len(scored)
    out: List[Dict] = []
    for rank, e in enumerate(scored):
        tier = ("Primary" if rank < max(1, n * 0.2)
                else "Secondary" if rank < max(2, n * 0.55) else "Supporting")
        out.append({"Entity": e["text"], "Type": e["type"],
                    "Description": spacy.explain(e["type"]) or e["type"],
                    "Tier": tier, "Importance": e["imp"],
                    "Occurrences": e["count"], "Context": e["desc"]})
    return out

# ── Structured deterministic extractors (regex + validating libraries) ──────────
# High-precision, deterministic, offline. Each candidate is validated/normalised
# by a purpose-built library so a greedy pattern can stay high-recall without
# bleeding false positives (e.g. phonenumbers validates, tldextract parses hosts).
import phonenumbers
import tldextract
import dateparser
import validators as _validators
from quantulum3 import parser as _qparser
from price_parser import Price as _Price

# Offline TLD parser: bundled Public Suffix List snapshot, never hits the network.
_TLD = tldextract.TLDExtract(suffix_list_urls=(), cache_dir=None)

_EMAIL_RE    = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_URL_RE      = re.compile(r"\b(?:https?://|www\.)[^\s<>()\[\]{}\"']+", re.I)
_PCT_RE      = re.compile(r"(?<![\w.])\d+(?:\.\d+)?\s?%")
_HASHTAG_RE  = re.compile(r"(?<!\w)#([A-Za-z]\w{1,49})")
_HANDLE_RE   = re.compile(r"(?<![\w@./])@([A-Za-z0-9_]{2,30})\b")
_DOI_RE      = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Za-z0-9]+", re.I)
_ARXIV_RE    = re.compile(r"\barXiv:\s?\d{4}\.\d{4,5}(?:v\d+)?\b", re.I)
_ISBN_RE     = re.compile(r"\bISBN(?:-1[03])?:?\s?(?:97[89][- ]?)?(?:\d[- ]?){9}[\dXx]\b")
_STAT_RE     = re.compile(
    r"(?:p\s?[<>=]\s?0?\.\d+|\bn\s?=\s?\d{1,9}\b|\br\s?=\s?-?0?\.\d+|"
    r"95%\s?CI|\bSD\s?=\s?[\d.]+|\bM\s?=\s?[\d.]+|t\s?\(\d+\)\s?=\s?-?[\d.]+)", re.I)
_CURRENCY_RE = re.compile(
    r"(?:[$€£¥₹]\s?\d[\d,]*(?:\.\d+)?\s?(?:[KMB]\b|million|billion|thousand)?"
    r"|\b\d[\d,]*(?:\.\d+)?\s?(?:USD|EUR|GBP|JPY|INR|CAD|AUD)\b)", re.I)
_DURATION_RE = re.compile(
    r"\b\d+(?:\.\d+)?[\-\s]?(?:seconds?|secs?|minutes?|mins?|hours?|hrs?|"
    r"days?|weeks?|months?|years?|milliseconds?)\b", re.I)
# Acronym definition: "Full Term Here (FTH)" — keep only when initials roughly match.
_ACRONYM_RE  = re.compile(
    r"\b((?:[A-Z][A-Za-z0-9]*\W+){1,6}[A-Za-z0-9]+)\s+\(([A-Z][A-Z0-9]{1,6})s?\)")
_DATE_RE     = re.compile(
    r"\b(?:\d{4}-\d{2}-\d{2}"                                            # ISO 2024-03-14
    r"|\d{1,2}[/-]\d{1,2}[/-]\d{2,4}"                                    # 03/14/2024
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}"
    r"|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}"
    r"|Q[1-4]\s+\d{4})\b", re.I)

def _ctx(text: str, start: int, end: int, pad: int = 55) -> str:
    """Short surrounding snippet for a matched span (for the Context column)."""
    return _clean(text[max(0, start - pad):min(len(text), end + pad)], 140)

def _acronym_phrase(initials: str, phrase: str) -> Optional[str]:
    """Return the trailing sub-phrase whose word-initials spell `initials` exactly
    (so 'degrees C. The Service Level Agreement (SLA)' → 'Service Level Agreement')."""
    words = re.findall(r"[A-Za-z0-9&]+", phrase)
    k = len(initials)
    if len(words) < k:
        return None
    tail = words[-k:]
    if "".join(w[0] for w in tail).upper() == initials.upper():
        return " ".join(tail)
    return None

_CURR_MULT = [("thousand", 1e3), ("million", 1e6), ("billion", 1e9),
              ("k", 1e3), ("m", 1e6), ("b", 1e9)]

def _currency_detail(v: str) -> str:
    """Normalise a currency surface form to '<amount> <symbol>', expanding K/M/B."""
    p = _Price.fromstring(v)
    if p.amount is None:
        return ""
    amt, low = p.amount_float, v.lower()
    for sfx, factor in _CURR_MULT:
        if re.search(rf"\d\s*{sfx}\b", low):
            amt *= factor
            break
    cur = p.currency or ""
    return (f"{amt:,.0f} {cur}".strip() if amt >= 1000 else f"{amt:g} {cur}".strip())

def extract_structured(text: str) -> List[Dict]:
    """L2: emails, URLs/domains, phones, dates, durations, measurements, currency,
    percentages, hashtags, handles, citations, statistics, acronyms — one flat,
    validated, deduplicated table of {Type, Value, Detail, Context}."""
    rows: List[Dict] = []
    seen: set = set()

    def add(typ: str, value: str, detail: str, ctx: str):
        key = (typ, value.lower().strip(), detail.lower().strip())
        if value.strip() and key not in seen:
            seen.add(key)
            rows.append({"Type": typ, "Value": value.strip(),
                         "Detail": detail, "Context": ctx})

    # Emails (validated)
    for m in _EMAIL_RE.finditer(text):
        v = m.group(0)
        if _validators.email(v):
            add("Email", v, v.split("@", 1)[1], _ctx(text, *m.span()))

    # URLs / domains (offline host parse → subdomain.domain.tld)
    for m in _URL_RE.finditer(text):
        v = m.group(0).rstrip(".,);:'\"")
        ex = _TLD(v)
        if ex.suffix:
            host = ".".join(p for p in (ex.subdomain, ex.domain, ex.suffix) if p)
            add("URL", v, host, _ctx(text, *m.span()))

    # Phone numbers (libphonenumber — validated, normalised to E.164)
    for region in ("US", None):
        try:
            for m in phonenumbers.PhoneNumberMatcher(text, region):
                if phonenumbers.is_valid_number(m.number):
                    e164 = phonenumbers.format_number(
                        m.number, phonenumbers.PhoneNumberFormat.E164)
                    add("Phone", m.raw_string, e164, _ctx(text, m.start, m.end))
        except Exception:
            pass

    # Dates (regex find → dateparser normalise to ISO)
    for m in _DATE_RE.finditer(text):
        v = m.group(0)
        try:
            dt = dateparser.parse(v)
        except Exception:
            dt = None
        add("Date", v, dt.strftime("%Y-%m-%d") if dt else "", _ctx(text, *m.span()))

    # Durations
    for m in _DURATION_RE.finditer(text):
        add("Duration", m.group(0), "", _ctx(text, *m.span()))

    # Measurements / quantities (quantulum3 — unit-bearing only; drop bare numbers)
    try:
        for q in _qparser.parse(text):
            unit = q.unit.name
            if unit and unit != "dimensionless":
                add("Measurement", q.surface, f"{q.value} {unit}", "")
    except Exception:
        pass

    # Currency (regex find → price-parser normalise amount + currency, expand K/M/B)
    for m in _CURRENCY_RE.finditer(text):
        v = m.group(0)
        add("Currency", v, _currency_detail(v), _ctx(text, *m.span()))

    # Percentages
    for m in _PCT_RE.finditer(text):
        add("Percentage", m.group(0).replace(" ", ""), "", _ctx(text, *m.span()))

    # Hashtags / social handles
    for m in _HASHTAG_RE.finditer(text):
        add("Hashtag", "#" + m.group(1), "", _ctx(text, *m.span()))
    for m in _HANDLE_RE.finditer(text):
        add("Social Handle", "@" + m.group(1), "", _ctx(text, *m.span()))

    # Citations (DOI / arXiv / ISBN)
    for rx, lbl in ((_DOI_RE, "DOI"), (_ARXIV_RE, "arXiv"), (_ISBN_RE, "ISBN")):
        for m in rx.finditer(text):
            add("Citation", m.group(0), lbl, _ctx(text, *m.span()))

    # Statistics (p-values, n=, r=, CI, SD, M, t-tests)
    for m in _STAT_RE.finditer(text):
        add("Statistic", m.group(0), "", _ctx(text, *m.span()))

    # Acronyms + expansion (Schwartz-Hearst-style, initials must trace the phrase)
    for m in _ACRONYM_RE.finditer(text):
        exp = _acronym_phrase(m.group(2), m.group(1))
        if exp:
            add("Acronym", m.group(2), exp, _ctx(text, *m.span()))

    type_order = {t: i for i, t in enumerate(
        ["Email", "Phone", "URL", "Social Handle", "Hashtag", "Date", "Duration",
         "Measurement", "Currency", "Percentage", "Statistic", "Citation",
         "Acronym"])}
    rows.sort(key=lambda r: (type_order.get(r["Type"], 99), r["Value"].lower()))
    return rows

# ── Zero-shot typed entities (GLiNER) ───────────────────────────────────────────
def extract_typed_entities(text: str) -> List[Dict]:
    """L4: GLiNER zero-shot NER for GLINER_LABELS (job title / product / event).
    Chunked + batched; deduplicated; keeps the highest-confidence hit per entity."""
    if GLINER is None:
        return []
    chunks = chunk_text(text, 1_200)            # GLiNER prefers short windows
    best: Dict[Tuple[str, str], Dict] = {}
    try:
        preds = GLINER.batch_predict_entities(
            chunks, GLINER_LABELS, threshold=GLINER_THRESHOLD)
    except Exception:
        preds = []
        for c in chunks:
            try: preds.append(GLINER.predict_entities(c, GLINER_LABELS,
                                                      threshold=GLINER_THRESHOLD))
            except Exception: preds.append([])
    for chunk, ents in zip(chunks, preds):
        for e in ents:
            name = _clean(e["text"], 80)
            if len(name) < 2:
                continue
            key = (name.lower(), e["label"])
            score = round(float(e["score"]), 3)
            if key not in best or score > best[key]["Confidence"]:
                s, en = e.get("start", 0), e.get("end", 0)
                best[key] = {"Entity": name, "Label": e["label"], "Confidence": score,
                             "Context": _ctx(chunk, s, en)}
    out = list(best.values())
    out.sort(key=lambda r: (r["Label"], -r["Confidence"]))
    return out

# ── Variables (regex) ───────────────────────────────────────────────────────────
def extract_variables(text: str) -> List[Dict]:
    results, seen = [], set()
    for sent in sent_tokenize(text):
        for pat, vtype in VARIABLE_PATTERNS:
            for m in re.finditer(pat, sent, re.IGNORECASE):
                try:
                    g = m.groups()
                    name, val = (g[0].strip(), g[1].strip()) if len(g) >= 2 else (None, None)
                    if not name or not val or len(name) < 2: continue
                    key = (name.lower()[:30], val[:20])
                    if key not in seen:
                        seen.add(key)
                        results.append({"Variable / Parameter": name, "Value": val,
                                        "Type": vtype, "Context": _clean(sent)})
                except Exception: continue
    return results

# ── Causal / relational regex extraction (with confidence) ──────────────────────
def extract_causal(text: str) -> List[Dict]:
    results, seen = [], set()
    for sent in sent_tokenize(text):
        for pat, rel in CAUSAL_PATTERNS:
            for m in re.finditer(pat, sent, re.IGNORECASE):
                try:
                    s, o = m.group(1).strip(), m.group(2).strip()
                    key = (s.lower()[:40], rel, o.lower()[:40])
                    if key not in seen and len(s) > 2 and len(o) > 2:
                        seen.add(key)
                        results.append({"Subject": s, "Relationship": rel, "Object": o,
                                        "Rel. Type": "causal",
                                        "Confidence": _CAUSAL_CONF.get(rel, 65),
                                        "Source Sentence": _clean(sent, 250)})
                except Exception: continue
    return results[:200]

def _display_rel(r: Dict) -> Dict:
    """Strip internal keys (_si …) and fix column order for the relationships table."""
    return {"Subject": r.get("Subject", ""), "Relationship": r.get("Relationship", ""),
            "Object": r.get("Object", ""), "Rel. Type": r.get("Rel. Type", ""),
            "Confidence": r.get("Confidence", ""), "Source Sentence": r.get("Source Sentence", "")}

# ── Dependency mapping ──────────────────────────────────────────────────────────
def build_dependencies(rels: List[Dict]) -> List[Dict]:
    G = nx.DiGraph()
    for r in rels:
        rel = r.get("Relationship", "").lower()
        if not any(w == rel or (" " in w and w in rel) for w in DEP_REL_WORDS):
            continue
        s = _node_label(r.get("Subject", "")).lower(); o = _node_label(r.get("Object", "")).lower()
        if not s or not o or s == o: continue
        if G.has_edge(s, o): G[s][o]["w"] += 1
        else: G.add_edge(s, o, w=1)
    if G.number_of_nodes() == 0:
        return []

    cyc_nodes = set()
    try:
        for c in nx.simple_cycles(G): cyc_nodes.update(c)
    except Exception:
        pass

    out = []
    for node in G.nodes():
        direct = sorted(G.successors(node))
        if not direct: continue
        try: total = len(nx.descendants(G, node))
        except Exception: total = len(direct)
        out.append({"Node": node, "Depends On": ", ".join(direct[:6]),
                    "Direct Deps": len(direct), "Total (incl. indirect)": total,
                    "Strength": sum(G[node][d]["w"] for d in direct),
                    "Circular?": "Yes" if node in cyc_nodes else "No"})
    out.sort(key=lambda d: (-d["Total (incl. indirect)"], -d["Strength"]))
    return out[:80]

# ── Causal-chain construction ───────────────────────────────────────────────────
def build_causal_chains(causal: List[Dict]) -> List[Dict]:
    G = nx.DiGraph(); conf: Dict[Tuple[str, str], float] = {}
    for r in causal:
        s = _node_label(r.get("Subject", "")).lower(); o = _node_label(r.get("Object", "")).lower()
        if not s or not o or s == o: continue
        G.add_edge(s, o); conf[(s, o)] = max(conf.get((s, o), 0), r.get("Confidence", 65))
    if G.number_of_edges() == 0:
        return []

    # Break cycles so longest-path search terminates.
    try:
        while True:
            cyc = nx.find_cycle(G, orientation="original")
            G.remove_edge(cyc[0][0], cyc[0][1])
    except nx.NetworkXNoCycle:
        pass

    roots = [n for n in G.nodes if G.in_degree(n) == 0]
    chains, seen, guard = [], set(), 0
    for root in roots:
        stack = [(root, [root])]; best = None
        while stack and guard < 50_000:
            guard += 1
            node, path = stack.pop()
            succ = [s for s in G.successors(node) if s not in path]
            if not succ:
                if best is None or len(path) > len(best): best = path
            for s in succ:
                stack.append((s, path + [s]))
        if best and len(best) >= 2:
            key = tuple(best)
            if key in seen: continue
            seen.add(key)
            cs = [conf.get((best[k], best[k + 1]), 65) for k in range(len(best) - 1)]
            chains.append({"Causal Chain": " → ".join(best), "Length": len(best),
                           "Root Cause": best[0], "Final Outcome": best[-1],
                           "Avg Confidence": round(sum(cs) / len(cs), 1)})
    chains.sort(key=lambda c: (-c["Length"], -c["Avg Confidence"]))
    return chains[:60]

# ── Contradiction detection ─────────────────────────────────────────────────────
def detect_contradictions(rels: List[Dict], variables: List[Dict],
                          sents: List[str]) -> List[Dict]:
    out, seen = [], set()

    # 1 ── same variable assigned different values
    byname: Dict[str, list] = defaultdict(list)
    for v in variables:
        vals = byname[v["Variable / Parameter"].lower()]
        if v["Value"] not in vals:
            vals.append(v["Value"])
    for name, vals in byname.items():
        if len(vals) > 1:
            out.append({"Statement A": f"{name} = {vals[0]}",
                        "Statement B": f"{name} = {vals[1]}",
                        "Conflict Type": "variable value conflict", "Severity": "Medium"})

    # 2 ── affirmation/negation + opposing-predicate conflicts on the same pair
    anti: Dict[str, str] = {}
    for a, b in ANTONYM_PAIRS: anti[a] = b; anti[b] = a
    pair_map: Dict[Tuple[str, str], list] = defaultdict(list)
    for r in rels:
        s = r.get("Subject", "").lower().strip(); o = r.get("Object", "").lower().strip()
        if len(s) < 3 or len(o) < 3: continue
        src = r.get("Source Sentence", "")
        neg = any(ng in src.lower() for ng in _NEGATIONS)
        pair_map[(s, o)].append((r.get("Relationship", "").lower(), neg, src))
    for (s, o), lst in pair_map.items():
        for i in range(len(lst)):
            for j in range(i + 1, len(lst)):
                r1, n1, src1 = lst[i]; r2, n2, src2 = lst[j]
                if r1 == r2 and n1 != n2:
                    ctype = "affirmation vs negation"
                elif anti.get(r1) == r2 or anti.get(r2) == r1:
                    ctype = "opposing relationship"
                else:
                    continue
                key = (s, o, ctype)
                if key in seen: continue
                seen.add(key)
                out.append({"Statement A": _clean(src1, 160) or f"{s} {r1} {o}",
                            "Statement B": _clean(src2, 160) or f"{s} {r2} {o}",
                            "Conflict Type": ctype, "Severity": "High"})
    return out[:60]

# ── Assumption discovery ────────────────────────────────────────────────────────
def detect_assumptions(sents: List[str]) -> List[Dict]:
    out, seen = [], set()
    for s in sents:
        sl = s.lower()
        for marker, conf in ASSUMPTION_MARKERS.items():
            if marker in sl:
                key = sl[:80]
                if key in seen: break
                seen.add(key)
                explicit = conf >= 70
                out.append({"Assumption": _clean(s, 220),
                            "Type": "Explicit" if explicit else "Implicit",
                            "Marker": marker, "Confidence": conf,
                            "Impact if False": ("Core conclusions built on this premise may fail."
                                                if explicit else
                                                "Some downstream reasoning may be weakened.")})
                break
    out.sort(key=lambda a: -a["Confidence"])
    return out[:60]

# ── Problem detection (severity tiers + root causes + affected entities) ─────────
def identify_problems(rels: List[Dict], entities: List[Dict], text: str,
                      sents: List[str], causal: List[Dict]) -> List[Dict]:
    cg = nx.DiGraph()
    for r in causal:
        s = _node_label(r.get("Subject", "")).lower(); o = _node_label(r.get("Object", "")).lower()
        if s and o and s != o: cg.add_edge(s, o)

    G = nx.DiGraph()
    for r in rels:
        s, o = r.get("Subject", ""), r.get("Object", "")
        if s and o: G.add_edge(s.lower()[:40], o.lower()[:40])
    cycles = []
    try: cycles = list(nx.simple_cycles(G))
    except Exception: pass
    cycle_nodes = {n for c in cycles for n in c}

    ent_names = [e["Entity"] for e in entities]
    problems, seen = [], set()
    for r in rels:
        s = r.get("Subject", "").lower(); rel = r.get("Relationship", "").lower()
        o = r.get("Object", "").lower(); src = r.get("Source Sentence", "")
        srcl = src.lower()
        # Match on whole words, not substrings — otherwise "issue" fires inside
        # "tissue" and "exploit" inside "exploitation", flagging ordinary prose
        # as Critical security risks. Keyword sets are all single words.
        src_w = _word_set(srcl); rel_w = _word_set(rel)
        triple_w = rel_w | _word_set(s) | _word_set(o)

        if (SECURITY_WORDS & src_w) or (SECURITY_WORDS & triple_w):
            sev, ptype = "Critical", "security / safety risk"
        elif PROBLEM_WORDS & triple_w:
            sev, ptype = "High", "problem keyword"
        elif s[:40] in cycle_nodes or o[:40] in cycle_nodes:
            sev, ptype = "Medium", "circular dependency"
        elif PROBLEM_WORDS & src_w:
            sev, ptype = "Low", "problematic context"
        else:
            continue

        key = (s[:40], rel[:20], o[:40])
        if key in seen: continue
        seen.add(key)

        nl = _node_label(r.get("Subject", "")).lower()
        roots = sorted(nx.ancestors(cg, nl))[:3] if cg.has_node(nl) else []
        affected = [e for e in ent_names if e.lower() in srcl][:4]
        problems.append({"Subject": r["Subject"], "Relationship": r["Relationship"],
                         "Object": r["Object"], "Problem Type": ptype, "Severity": sev,
                         "Root Causes": ", ".join(roots),
                         "Affected Entities": ", ".join(affected),
                         "Source Sentence": _clean(src, 200)})

    sev_rank = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    problems.sort(key=lambda p: sev_rank.get(p["Severity"], 9))
    return problems

# ── Knowledge-graph connectivity helpers (entity standardisation + inference) ────
# Relations that are logically transitive (A r B ∧ B r C ⇒ A r C). Restricting
# inference to these keeps inferred edges sound (unlike e.g. "prevents"). Maps
# both SVO verb-lemma forms ("depend") and causal-phrase forms ("depends on") to
# one canonical label so a chain spanning both still chains.
_TRANSITIVE_RELS = {
    "depend": "depends on", "depends": "depends on", "depends on": "depends on",
    "require": "requires", "requires": "requires",
    "need": "needs", "needs": "needs",
    "contain": "contains", "contains": "contains",
    "include": "contains", "includes": "contains",
    "part of": "part of",
    "lead to": "leads to", "leads to": "leads to",
    "result in": "results in", "results in": "results in",
    "cause": "causes", "causes": "causes",
}

def _standardize_nodes(keys: List[str], sim_threshold: float = 0.85) -> Dict[str, str]:
    """Merge synonymous / variant node keys into one canonical key so the graph
    stops fragmenting on surface form ("apple" vs "apple inc", "neural network"
    vs "neural net"). Uses KeyBERT's embedder (cosine) + whole-word subset.
    Deterministic and offline. Returns {key: canonical_key}."""
    uniq = sorted({k for k in keys if k})
    if len(uniq) < 2:
        return {k: k for k in uniq}
    freq = Counter(keys)
    sim = cosine_similarity(_embed(uniq))
    parent = list(range(len(uniq)))
    def find(i):
        while parent[i] != i:
            parent[i] = parent[parent[i]]; i = parent[i]
        return i
    toks = [set(u.split()) for u in uniq]
    for i in range(len(uniq)):
        for j in range(i + 1, len(uniq)):
            subset = toks[i] and toks[j] and (toks[i] <= toks[j] or toks[j] <= toks[i])
            if sim[i][j] >= sim_threshold or subset:
                parent[find(i)] = find(j)
    groups: Dict[int, List[str]] = defaultdict(list)
    for i in range(len(uniq)):
        groups[find(i)].append(uniq[i])
    cmap: Dict[str, str] = {}
    for members in groups.values():                    # canonical = most frequent
        rep = sorted(members, key=lambda k: (-freq[k], -len(k), k))[0]
        for m in members:
            cmap[m] = rep
    return cmap

def _infer_relationships(G: nx.DiGraph, emb_of: Dict[str, np.ndarray],
                         max_transitive: int = 40, sim_threshold: float = 0.55) -> int:
    """Densify the graph without an LLM: (1) transitive closure over transitive-
    safe relations, (2) bridge disconnected components by embedding similarity.
    Mutates G (adds edges flagged inferred=True). Returns count of edges added."""
    added = 0
    # (1) transitive: A —r→ B —r→ C  ⇒  A —r→ C  (same canonical transitive-safe relation)
    for a in sorted(G.nodes()):
        for b in sorted(G.successors(a)):
            rel = _TRANSITIVE_RELS.get((G[a][b].get("label") or "").lower())
            if rel is None:
                continue
            for c in sorted(G.successors(b)):
                if c == a or G.has_edge(a, c): continue
                if _TRANSITIVE_RELS.get((G[b][c].get("label") or "").lower()) != rel: continue
                G.add_edge(a, c, label=rel, rtype="inferred (transitive)",
                           sentence=f"Inferred: {a} {rel} {b}, {b} {rel} {c}",
                           problem=False, weight=1, inferred=True)
                added += 1
                if added >= max_transitive: break
            if added >= max_transitive: break
        if added >= max_transitive: break

    # (2) bridge components: attach each satellite component to the main one via
    #     its single most-similar node pair (single-linkage across components).
    comps = sorted(nx.connected_components(G.to_undirected()), key=len, reverse=True)
    if len(comps) > 1:
        main = [n for n in comps[0] if n in emb_of]
        for comp in comps[1:]:
            best = None
            for n in (x for x in comp if x in emb_of):
                for m in main:
                    s = float(cosine_similarity([emb_of[n]], [emb_of[m]])[0][0])
                    if best is None or s > best[0]:
                        best = (s, n, m)
            if best and best[0] >= sim_threshold:
                _, n, m = best
                G.add_edge(n, m, label="related to", rtype="inferred (similarity)",
                           sentence=f"Inferred link (cosine {best[0]:.2f})",
                           problem=False, weight=1, inferred=True)
                main.extend(x for x in comp if x in emb_of)   # allow later chaining
                added += 1
    return added

# ── Knowledge graph (influence / betweenness / key nodes) ───────────────────────
def build_knowledge_graph(rels: List[Dict], entities: List[Dict],
                          problems: List[Dict], concepts: List[Dict],
                          max_nodes: int = 120) -> Dict:
    ent_type, ent_count = {}, {}
    for e in entities:
        k = e["Entity"].lower()
        if k not in ent_type:
            ent_type[k] = e["Type"]; ent_count[k] = e.get("Occurrences", 1)

    # 1 ── canonical surfaces (strip leading articles) → merge synonyms/variants
    triples: List[Tuple[str, str, Dict]] = []
    surface_of: Dict[str, str] = {}
    for r in rels:
        ss = _node_label(r.get("Subject", ""), 40); oo = _node_label(r.get("Object", ""), 40)
        s, o = ss.lower(), oo.lower()
        if not s or not o or s == o: continue
        # Central guard: a bare pronoun/relativiser must never become a node,
        # whatever extractor produced it (SVO, causal, …). Keeps "they"/"it"/
        # "that" from collapsing many sentences onto one meaningless hub.
        if s in _PRONOUN_LEMMAS or o in _PRONOUN_LEMMAS: continue
        triples.append((s, o, r))
        surface_of.setdefault(s, ss); surface_of.setdefault(o, oo)
    if not triples:
        return {"nodes": [], "edges": [],
                "stats": {"nodes": 0, "edges": 0, "communities": 0, "inferred_edges": 0}}

    cmap = _standardize_nodes([t[0] for t in triples] + [t[1] for t in triples])
    def canon(label: str) -> str:
        k = _node_label(label, 40).lower()
        return cmap.get(k, k)
    problem_pairs = {(canon(p["Subject"]), canon(p["Object"])) for p in problems}

    G = nx.DiGraph()
    for s, o, r in triples:
        sk, ok = cmap.get(s, s), cmap.get(o, o)
        if sk == ok: continue
        for nk in (sk, ok):
            if G.has_node(nk): G.nodes[nk]["weight"] += 1
            else: G.add_node(nk, label=surface_of.get(nk, nk),
                             etype=ent_type.get(nk, ""), weight=max(ent_count.get(nk, 1), 1))
        if G.has_edge(sk, ok): G[sk][ok]["weight"] += 1
        else: G.add_edge(sk, ok, label=r.get("Relationship", ""),
                         rtype=r.get("Rel. Type", ""),
                         sentence=_clean(r.get("Source Sentence", ""), 200),
                         problem=(sk, ok) in problem_pairs, weight=1, inferred=False)

    if G.number_of_nodes() > max_nodes:
        keep = [n for n, _ in sorted(G.degree, key=lambda x: -x[1])[:max_nodes]]
        G = G.subgraph(keep).copy()

    # 2 ── relationship inference (transitive + similarity bridging) to densify
    inferred = 0
    if G.number_of_nodes() > 2:
        node_keys = sorted(G.nodes())
        vecs = _embed([surface_of.get(n, n) for n in node_keys])
        emb_of = {n: vecs[i] for i, n in enumerate(node_keys)}
        inferred = _infer_relationships(G, emb_of)
    G.remove_nodes_from(list(nx.isolates(G)))

    comm_of: Dict[str, int] = {}
    try:
        if G.number_of_nodes() > 2:
            comms = nx.community.greedy_modularity_communities(G.to_undirected())
            comm_of = {n: i for i, c in enumerate(comms) for n in c}
    except Exception:
        pass

    deg_c = nx.degree_centrality(G) if G.number_of_nodes() else {}
    try:
        btw_c = nx.betweenness_centrality(G) if G.number_of_nodes() > 2 else {}
    except Exception:
        btw_c = {}

    def category(nk: str) -> str:
        if ent_type.get(nk): return ent_type[nk]
        return "Concept"

    influence = {n: deg_c.get(n, 0.0) * G.nodes[n].get("weight", 1) for n in G.nodes()}
    nodes = [{"id": n, "label": d["label"], "type": category(n),
              "size": d.get("weight", 1), "centrality": round(deg_c.get(n, 0.0), 4),
              "betweenness": round(btw_c.get(n, 0.0), 4),
              "influence": round(influence.get(n, 0.0), 4),
              "group": comm_of.get(n, 0)}
             for n, d in G.nodes(data=True)]
    edges = [{"from": u, "to": v, "label": d.get("label", ""), "rtype": d.get("rtype", ""),
              "sentence": d.get("sentence", ""), "problem": bool(d.get("problem")),
              "inferred": bool(d.get("inferred")), "weight": d.get("weight", 1)}
             for u, v, d in G.edges(data=True)]

    key: Dict[str, str] = {}
    if nodes:
        key["most_influential"] = max(nodes, key=lambda n: n["influence"])["label"]
        key["most_connected"] = max(nodes, key=lambda n: n["centrality"])["label"]
        if btw_c:
            key["bridge"] = max(nodes, key=lambda n: n["betweenness"])["label"]
    return {"nodes": nodes, "edges": edges,
            "stats": {"nodes": len(nodes), "edges": len(edges),
                      "communities": len(set(comm_of.values())) if comm_of else 0,
                      "inferred_edges": inferred, **key}}

# ── Executive intelligence roll-up ──────────────────────────────────────────────
def build_executive_insights(entities, concepts, variables, problems, rels,
                              dependencies, causal_chains, contradictions,
                              assumptions, graph) -> List[Dict]:
    ins: List[Dict] = []
    def add(cat, finding): ins.append({"Category": cat, "Finding": _clean(str(finding), 300)})

    if concepts and concepts[0].get("Concept"):
        add("Most Important Concept",
            f'{concepts[0]["Concept"]} (importance {concepts[0]["Importance"]})')
    if entities:
        add("Primary Entity",
            f'{entities[0]["Entity"]} — {entities[0]["Type"]}, {entities[0]["Occurrences"]} mentions')

    crit = [p for p in problems if p["Severity"] in ("Critical", "High")]
    if crit:
        add(f'Most Critical Problem ({crit[0]["Severity"]})',
            f'{crit[0]["Subject"]} {crit[0]["Relationship"]} {crit[0]["Object"]}')

    gk = graph.get("stats", {})
    if gk.get("most_influential"): add("Most Influential Node", gk["most_influential"])
    if gk.get("bridge"): add("Key Bridge / Connector", gk["bridge"])

    themes = Counter(c.get("Theme / Cluster", "") for c in concepts if c.get("Theme / Cluster"))
    if themes:
        t, cnt = themes.most_common(1)[0]
        add("Dominant Theme", f'"{t}" anchors {cnt} related concepts')

    circ = [d for d in dependencies if d.get("Circular?") == "Yes"]
    if circ:
        add("System Weakness", f'{len(circ)} circular dependencies (e.g. {circ[0]["Node"]})')
    if contradictions:
        add("Risk — Contradiction",
            f'{len(contradictions)} conflicting statements (e.g. {contradictions[0]["Conflict Type"]})')
    strong_assum = [a for a in assumptions if a["Confidence"] >= 78]
    if strong_assum:
        add("Risk — Assumption",
            f'{len(strong_assum)} strong assumptions underpin the reasoning')

    if causal_chains:
        rc = Counter(c["Root Cause"] for c in causal_chains).most_common(1)[0][0]
        add("Root Cause", f'"{rc}" initiates the longest causal chain')

    if crit: add("Recommendation", f'Prioritise mitigation of {len(crit)} high/critical problems')
    if circ: add("Recommendation", "Break circular dependencies to improve robustness")
    if contradictions: add("Recommendation", "Reconcile contradictory statements before acting")
    if not crit and not contradictions:
        add("Strategic Insight", "No critical problems or contradictions — content is internally consistent")
    return ins

def _structured_insights(structured: List[Dict], typed: List[Dict]) -> List[Dict]:
    """Fold the deterministic L2 / typed-entity findings into the insight roll-up."""
    ins: List[Dict] = []
    by_type: Dict[str, list] = defaultdict(list)
    for r in structured:
        by_type[r["Type"]].append(r["Value"])
    by_label: Dict[str, list] = defaultdict(list)
    for r in typed:
        by_label[r["Label"]].append(r["Entity"])

    def add(cat, finding): ins.append({"Category": cat, "Finding": _clean(str(finding), 300)})

    contacts = len(by_type.get("Email", [])) + len(by_type.get("Phone", []))
    if contacts:
        sample = (by_type.get("Email") or by_type.get("Phone"))[0]
        add("Contact Points", f"{contacts} contact(s) detected (e.g. {sample})")
    if by_type.get("Currency"):
        add("Financial Figures",
            f'{len(by_type["Currency"])} monetary amounts (e.g. {by_type["Currency"][0]})')
    if by_type.get("Statistic"):
        add("Statistical Evidence",
            f'{len(by_type["Statistic"])} statistics reported (e.g. {by_type["Statistic"][0]})')
    if by_type.get("Date"):
        add("Temporal Span", f'{len(by_type["Date"])} dates referenced')
    if by_type.get("Acronym"):
        add("Terminology", f'{len(by_type["Acronym"])} acronyms defined in-text')
    for label in ("job title", "product", "event"):
        if by_label.get(label):
            add(f"Typed Entities — {label.title()}",
                f'{len(by_label[label])} found (e.g. {by_label[label][0]})')
    return ins

# ── Export helpers ────────────────────────────────────────────────────────────
def _export_excel_bytes(results: Dict) -> bytes:
    wb = Workbook()
    HF = Font(bold=True, color="FFFFFF", size=11)
    HB = PatternFill("solid", fgColor="1F4E79")
    RH = PatternFill("solid", fgColor="B71C1C")
    RF = PatternFill("solid", fgColor="FCE4EC")
    YF = PatternFill("solid", fgColor="FFF9C4")

    def ws(name, data, hfill=None):
        sheet = wb.create_sheet(name)
        if not data: sheet.append(["No data."]); return sheet
        df = pd.DataFrame(data)
        sheet.append(list(df.columns))
        for c in sheet[1]:
            c.font = HF; c.fill = hfill or HB
            c.alignment = Alignment(horizontal="center", wrap_text=True)
        for row in df.itertuples(index=False): sheet.append(list(row))
        for col in sheet.columns:
            w = max((len(str(c.value or "")) for c in col), default=10)
            sheet.column_dimensions[col[0].column_letter].width = min(w + 2, 55)
        return sheet

    s0 = wb.active; s0.title = "Summary"
    s0.append(["Metric", "Value"])
    for c in s0[1]: c.font = HF; c.fill = HB; c.alignment = Alignment(horizontal="center")
    for k, v in results["summary"].items(): s0.append([k, v])
    for col in s0.columns: s0.column_dimensions[col[0].column_letter].width = 35

    ws("Files",                  results["files"])
    ws("Entities",               results["entities"])
    ws("Typed Entities",         results["typed_entities"])
    ws("Structured Data",        results["structured"])
    ws("Key Concepts",           results["concepts"])
    ws("Variables",              results["variables"])
    ws("Relationships",          results["relationships"])
    ws("Dependencies",           results["dependencies"])
    ws("Causal Chains",          results["causal_chains"])
    sp = ws("Problem Relationships", results["problems"], hfill=RH)
    if results["problems"]:
        keys = list(results["problems"][0].keys())
        sev_col = keys.index("Severity") + 1
        for i in range(2, sp.max_row + 1):
            fill = RF if sp.cell(i, sev_col).value in ("High", "Critical") else YF
            for cell in sp[i]: cell.fill = fill
    ws("Contradictions",         results["contradictions"], hfill=RH)
    ws("Assumptions",            results["assumptions"])
    ws("Executive Insights",     results["insights"])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()

def _export_csv_bytes(results: Dict) -> bytes:
    buf = io.BytesIO()
    sheets = {"summary": [results["summary"]], "files": results["files"],
              "entities": results["entities"], "typed_entities": results["typed_entities"],
              "structured": results["structured"], "concepts": results["concepts"],
              "variables": results["variables"], "relationships": results["relationships"],
              "dependencies": results["dependencies"], "causal_chains": results["causal_chains"],
              "problems": results["problems"], "contradictions": results["contradictions"],
              "assumptions": results["assumptions"], "insights": results["insights"]}
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in sheets.items():
            csv_str = pd.DataFrame(data if data else [{name: "No data"}]).to_csv(index=False, encoding="utf-8-sig")
            zf.writestr(f"{name}.csv", csv_str)
    return buf.getvalue()

def _export_markdown_bytes(results: Dict) -> bytes:
    def section(title, data):
        lines = [f"## {title}", ""]
        if not data: return lines + ["*No data.*", ""]
        df = pd.DataFrame(data)
        lines.append("| " + " | ".join(str(c) for c in df.columns) + " |")
        lines.append("|" + "|".join("---" for _ in df.columns) + "|")
        for _, row in df.iterrows():
            lines.append("| " + " | ".join(str(v)[:120].replace("|","\\|").replace("\n"," ") for v in row) + " |")
        return lines + [""]

    s = results["summary"]
    out = ["# Document Analysis Report", "", "## Summary", "",
           "| Metric | Value |", "|--------|-------|"]
    for k, v in s.items(): out.append(f"| {k} | {v} |")
    out += ["", "---", ""]
    out += section("Executive Insights",     results["insights"])
    out += section("Files Analyzed",         results["files"])
    out += section("Named Entities",         results["entities"])
    out += section("Typed Entities",         results["typed_entities"])
    out += section("Structured Data",        results["structured"])
    out += section("Key Concepts",           results["concepts"])
    out += section("Variables",              results["variables"])
    out += section("Relationships",          results["relationships"])
    out += section("Dependency Map",         results["dependencies"])
    out += section("Causal Chains",          results["causal_chains"])
    out += section("Problem Relationships",  results["problems"])
    out += section("Contradictions",         results["contradictions"])
    out += section("Assumptions",            results["assumptions"])
    return "\n".join(out).encode("utf-8")

# ── Flask app ─────────────────────────────────────────────────────────────────
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 64 * 1024 * 1024  # 64 MB

_RESULTS: Dict[str, dict] = {}  # session_id → results

@app.route("/")
def index():
    return render_template("index.html", nlp_model=NLP_MODEL)

@app.route("/analyze", methods=["POST"])
def analyze():
    uploaded = request.files.getlist("files")
    if not uploaded or all(f.filename == "" for f in uploaded):
        return jsonify({"error": "No files uploaded"}), 400

    records, texts = [], []
    for f in uploaded:
        name = f.filename
        suffix = Path(name).suffix.lower() or ".tmp"
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp_path = tmp.name
                f.save(tmp_path)
            text, ftype = read_document(tmp_path)
            # Counts describe the file as delivered; analysis runs on the body with
            # any trailing bibliography removed so it isn't mined as content.
            records.append({"Filename": name, "Format": ftype,
                             "Words": len(text.split()), "Characters": len(text), "Status": "OK"})
            texts.append(strip_references(text))
        except Exception as exc:
            records.append({"Filename": name, "Format": "?",
                             "Words": 0, "Characters": 0, "Status": f"Error: {exc}"})
        finally:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)

    combined = "\n\n".join(texts)
    if not combined.strip():
        return jsonify({"error": "No text could be extracted from the uploaded files."}), 400
    if len(combined) > 400_000:
        combined = combined[:400_000]

    # Preprocessing → deterministic L2 → single spaCy pass → KeyBERT + GLiNER → analyses.
    combined        = preprocess_text(combined)
    structured      = extract_structured(combined)
    sents, ent_counter, ent_sents, svo = parse_document(combined)
    concepts        = extract_concepts(combined, sents, ent_sents)
    entities        = build_entities(ent_counter, ent_sents, sents, concepts)
    typed_entities  = extract_typed_entities(combined)
    variables       = extract_variables(combined)
    causal          = extract_causal(combined)
    all_rels        = svo + causal
    problems        = identify_problems(all_rels, entities, combined, sents, causal)
    dependencies    = build_dependencies(all_rels)
    causal_chains   = build_causal_chains(causal)
    contradictions  = detect_contradictions(all_rels, variables, sents)
    assumptions     = detect_assumptions(sents)
    rels_display    = [_display_rel(r) for r in all_rels]
    graph           = build_knowledge_graph(all_rels, entities, problems, concepts)
    insights        = build_executive_insights(entities, concepts, variables, problems,
                                                all_rels, dependencies, causal_chains,
                                                contradictions, assumptions, graph)
    insights       += _structured_insights(structured, typed_entities)

    sid = str(uuid.uuid4())
    results = {
        "files": records, "entities": entities, "typed_entities": typed_entities,
        "concepts": concepts, "structured": structured,
        "variables": variables, "relationships": rels_display, "problems": problems,
        "dependencies": dependencies, "causal_chains": causal_chains,
        "contradictions": contradictions, "assumptions": assumptions,
        "insights": insights, "graph": graph,
        "summary": {
            "Files Analyzed":       len(records),
            "Total Words":          f"{sum(r.get('Words', 0) for r in records):,}",
            "Entities Found":       len(entities),
            "Typed Entities":       len(typed_entities),
            "Structured Data":      len(structured),
            "Key Concepts":         len(concepts),
            "Variables Detected":   len(variables),
            "Relationships Mapped": len(rels_display),
            "Dependencies":         len(dependencies),
            "Causal Chains":        len(causal_chains),
            "Problems":             len(problems),
            "Contradictions":       len(contradictions),
            "Assumptions":          len(assumptions),
            "Graph Nodes":          graph["stats"]["nodes"],
            "Graph Edges":          graph["stats"]["edges"],
            "NLP Model":            NLP_MODEL,
            "Concept Model":        "all-MiniLM-L6-v2 (KeyBERT)",
            "Typed-Entity Model":   "gliner_small-v2.1" if GLINER else "unavailable",
        },
    }
    _RESULTS[sid] = results
    return jsonify({"session_id": sid, "results": results})

@app.route("/export/<sid>/<fmt>")
def export_results(sid: str, fmt: str):
    results = _RESULTS.get(sid)
    if not results:
        return "Session not found or expired.", 404

    if fmt == "xlsx":
        data = _export_excel_bytes(results)
        return send_file(io.BytesIO(data), download_name="document_analysis.xlsx",
                         as_attachment=True,
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    if fmt == "csv":
        data = _export_csv_bytes(results)
        return send_file(io.BytesIO(data), download_name="document_analysis.zip",
                         as_attachment=True, mimetype="application/zip")
    if fmt == "md":
        data = _export_markdown_bytes(results)
        return send_file(io.BytesIO(data), download_name="document_analysis.md",
                         as_attachment=True, mimetype="text/markdown")
    return "Unknown format.", 400

def _lan_ip() -> str:
    """Pick the real LAN address, preferring private ranges over VPN/CGNAT (100.64/10)."""
    import socket
    candidates = []
    try:
        for info in socket.getaddrinfo(socket.gethostname(), None, socket.AF_INET):
            ip = info[4][0]
            if not ip.startswith(("127.", "169.254.")):
                candidates.append(ip)
    except Exception:
        pass
    for ip in candidates:
        if ip.startswith(("192.168.", "10.")) or \
           (ip.startswith("172.") and 16 <= int(ip.split(".")[1]) <= 31):
            return ip
    return candidates[0] if candidates else "127.0.0.1"

if __name__ == "__main__":
    import webbrowser, threading
    lan = _lan_ip()
    url = "http://localhost:5000"
    threading.Timer(1.2, lambda: webbrowser.open(url)).start()
    print(f"Data Analyzer Web running at:")
    print(f"  This computer : {url}")
    print(f"  On your WiFi  : http://{lan}:5000")
    app.run(host="0.0.0.0", port=5000, debug=True, use_reloader=False, threaded=True)

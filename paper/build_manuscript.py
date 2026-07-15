"""
Generates the technical white paper / manuscript draft (Technical_Whitepaper.docx)
describing the data-analyzer-web pipeline, formatted as a peer-review submission
draft (double-spaced body text, Times New Roman, numbered sections).

Run:  py build_manuscript.py
Output: Technical_Whitepaper.docx (same directory)

This script is a build tool for the manuscript, not part of the running app.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Base styling ──────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2.0
normal.paragraph_format.space_after = Pt(0)

for lvl, size in ((1, 14), (2, 13), (3, 12)):
    st = doc.styles[f'Heading {lvl}']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.italic = (lvl == 3)
    st.font.color.rgb = None
    st.paragraph_format.space_before = Pt(18)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.0
    # Word runs sometimes keep a leftover theme color; force automatic/black.
    rpr = st.element.get_or_add_rPr()
    color = OxmlElement('w:color'); color.set(qn('w:val'), '000000')
    rpr.append(color)

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text, numbered=True, counter=[0]):
    if numbered:
        counter[0] += 1
        text = f"{counter[0]}. {text}"
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def p(text, italic=False, center=False, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.italic = italic
    return para

def bullets(items):
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)

def numbered(items):
    for item in items:
        para = doc.add_paragraph(style='List Number')
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(4)
        para.add_run(item)

def table(headers, rows, widths=None, caption=None, cap_num=[0]):
    if caption:
        cap_num[0] += 1
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(10)
        cp.paragraph_format.line_spacing = 1.0
        r = cp.add_run(f"Table {cap_num[0]}. {caption}")
        r.bold = True; r.font.size = Pt(11)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True; run.font.size = Pt(10); run.font.name = 'Times New Roman'
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10); run.font.name = 'Times New Roman'
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def pagebreak():
    doc.add_page_break()

# ── Title page ────────────────────────────────────────────────────────────────
p("DocAnalyzer: A Modular, Fully Offline Pipeline for Structured Extraction,\n"
  "Contradiction Detection, and Hybrid-Retrieval Question Answering\n"
  "over Unstructured Documents",
  center=True, space_after=18)
doc.paragraphs[-1].runs[0].bold = True
doc.paragraphs[-1].runs[0].font.size = Pt(16)
doc.paragraphs[-1].paragraph_format.line_spacing = 1.2

p("[Author Name(s) — to be completed]", center=True, space_after=2)
p("[Affiliation — to be completed]", center=True, space_after=2)
p("[Corresponding author email — to be completed]", center=True, space_after=18)

p("Manuscript draft prepared for submission. Line spacing, citation style, "
  "figure placement, and page limits should be adjusted to the target "
  "journal's author guidelines prior to submission.", italic=True, center=True)

pagebreak()

# ── Abstract ──────────────────────────────────────────────────────────────────
h1("Abstract", numbered=False)
p(
"Long-form technical, scientific, and business documents routinely encode claims, "
"quantitative parameters, causal mechanisms, hidden assumptions, and — not "
"infrequently — internal contradictions that are difficult for a human reader to "
"audit exhaustively. We present DocAnalyzer, a modular, fully offline "
"document-understanding pipeline that converts an arbitrary input document (PDF, "
"Word, Markdown, or plain text) into eleven linked analytical layers: named "
"entities (with embedding-based re-typing), zero-shot typed entities, ranked key "
"concepts, extracted variables and quantities, confidence-scored subject–verb–"
"object relationships, dependency and causal-chain graphs, severity-tiered "
"problem/risk statements, multi-strategy contradiction detection (including "
"benefit–cost trade-off axes such as precision-versus-off-target or "
"speed-versus-safety), implicit assumption mining, and a canonicalized, "
"community-detected knowledge graph. Every layer above runs without any external "
"network call, combining classical dependency parsing (spaCy), zero-shot "
"named-entity typing (GLiNER), transformer sentence embeddings "
"(Sentence-BERT/KeyBERT), lexical scoring (BM25), and cross-encoder reranking for "
"a companion hybrid-retrieval question-answering interface; an optional cloud "
"large-language-model call is used only as a disclosed, non-required enrichment "
"step. We describe the system architecture and each extraction method in detail "
"and present a qualitative case study on a peer-reviewed genome-editing review "
"article paired with an introductory background text, illustrating the system's "
"output across all analytical layers. We discuss the pipeline's design "
"rationale and current limitations — most notably the absence of a labeled "
"benchmark corpus for formal precision/recall evaluation — and outline the "
"human-annotation study required before the extraction-accuracy claims in this "
"manuscript can be quantitatively substantiated."
)
p("Keywords: natural language processing; named entity recognition; zero-shot "
  "entity typing; open relation extraction; causal chain mining; contradiction "
  "detection; assumption mining; knowledge graph construction; hybrid retrieval; "
  "retrieval-augmented question answering; document understanding", italic=True)

pagebreak()

# ── 1. Introduction ─────────────────────────────────────────────────────────
h1("Introduction")
h2("1.1 Motivation")
p(
"Readers of technical reports, scientific reviews, regulatory filings, and "
"long-form business documents are routinely expected to identify not only what a "
"document states, but what it assumes, what quantitative claims it makes, how its "
"stated mechanisms causally chain together, and — critically — where it "
"contradicts itself or acknowledges a trade-off without resolving it. Manually "
"performed, this kind of audit does not scale: a single peer-reviewed review "
"article can contain hundreds of entity mentions, dozens of quantitative claims, "
"and causal chains spanning several intermediate steps, any of which may be "
"revised, retracted, or superseded elsewhere in the same document."
)
p(
"Large language models can perform many of these tasks directly, but at the cost "
"of opacity (the reasoning behind an extracted claim is not inspectable), "
"non-determinism (repeated runs may disagree), and a dependency on a paid, "
"rate-limited, internet-connected API — properties that are undesirable when the "
"goal is a reproducible, auditable analysis artifact suitable for institutional "
"or regulatory use. We instead ask: how much of this structured-extraction task "
"can be accomplished with a fully offline, deterministic pipeline built from "
"well-established NLP components — dependency parsing, zero-shot NER, "
"transformer sentence embeddings, lexical retrieval — with a large language "
"model used, if at all, only as an optional and disclosed supplementary step?"
)
h2("1.2 Contributions")
p("This paper makes the following contributions:")
bullets([
 "A fully offline, deterministic, multi-stage document-extraction pipeline that "
 "combines classical dependency parsing (spaCy), zero-shot typed named-entity "
 "recognition (GLiNER), transformer-embedding keyphrase ranking (KeyBERT), and a "
 "pattern-based library for quantitative variable extraction, all seeded from a "
 "single shared parse pass over the source document.",
 "A confidence-scored, per-triple relationship-extraction method (Section 4.6) "
 "and a graph-canonicalization procedure — embedding-based node merging plus a "
 "curated causal-verb whitelist — that allows a sparse set of directly labeled "
 "causal relations to be densified into multi-hop causal chains (Section 4.8).",
 "A contradiction-detection module composed of four complementary strategies: "
 "context-gated variable-value conflicts, negation/antonym relationship "
 "conflicts, discourse-marker contrastive-claim detection, and axis-based "
 "benefit–cost trade-off detection (e.g., precision-versus-off-target, "
 "speed-versus-safety) (Section 4.10).",
 "An implicit-assumption detector that combines an extensible lexical-marker "
 "table with a thematic co-occurrence heuristic to surface unstated premises "
 "about delivery, specificity, scalability, and safety (Section 4.11).",
 "A resolved, community-detected knowledge-graph construction procedure with "
 "salience-protected node pruning — so that low-degree but conceptually "
 "important nodes are not discarded by naive degree-based pruning — and "
 "transitive- and similarity-based edge inference (Section 4.12).",
 "A hybrid lexical (BM25) plus semantic (sentence-embedding) plus cross-encoder "
 "reranked retrieval interface for natural-language question answering over the "
 "extracted analysis, preceded by a battery of deterministic fast-path intent "
 "handlers for common query types (Section 4.14–4.15).",
 "A qualitative case study (Section 6) applying the full pipeline to a "
 "peer-reviewed genome-editing review article, illustrating system output across "
 "every analytical layer, together with an honest discussion of the evaluation "
 "methodology's current limitations (Section 7).",
])
h2("1.3 Paper Organization")
p(
"Section 2 situates this work relative to prior work in named entity recognition, "
"keyphrase extraction, open relation and causal extraction, contradiction and "
"stance detection, assumption mining, knowledge-graph construction, and "
"retrieval-augmented question answering. Section 3 gives a system-level overview "
"of the architecture. Section 4 details each extraction method. Section 5 "
"describes the implementation and deployment model. Section 6 presents the "
"qualitative case study. Section 7 discusses limitations, threats to validity, "
"and ethical considerations. Section 8 concludes and outlines future work."
)

# ── 2. Related Work ───────────────────────────────────────────────────────────
h1("Related Work")

h2("2.1 Named Entity Recognition and Zero-Shot Entity Typing")
p(
"Statistical and neural named entity recognition (NER) has long relied on "
"supervised sequence labeling over a fixed label inventory; spaCy's "
"convolutional/transition-based parser [1] is representative of this mature, "
"fast, and widely deployed approach and forms the backbone of our sentence "
"segmentation, dependency parsing, and base entity extraction. A complementary, "
"more recent line of work reframes NER as a zero-shot span-classification "
"problem so that arbitrary label sets can be supplied at inference time without "
"retraining; GLiNER [3] is one such generalist model and is used here to extract "
"typed entities (job titles, products, events) that fall outside spaCy's fixed "
"label inventory. Our pipeline further re-types spaCy's generic labels using "
"cosine similarity between sentence-embedded entity context and per-label "
"description corpora, an inexpensive alternative to fine-tuning a domain-specific "
"NER model."
)
h2("2.2 Keyphrase and Concept Extraction")
p(
"KeyBERT [2] extracts keyphrases by embedding candidate n-grams and the "
"document (or a local context window) with a Sentence-BERT-style encoder [4] "
"and ranking candidates by cosine similarity to the document embedding. We build "
"on this by adding an importance-ranking and percentile-based tiering layer "
"(Core / Supporting / Derived) and a lightweight thematic-clustering pass so "
"that the resulting concept list can be grouped and prioritized rather than "
"presented as an unranked bag of keyphrases."
)
h2("2.3 Open Relation Extraction and Causal Relation Mining")
p(
"Open information extraction typically derives subject-predicate-object triples "
"from a dependency parse, trading recall against the fixed-schema precision of "
"supervised relation extraction. We follow this dependency-parse-based approach "
"for our general relationship layer, and additionally maintain a separate, "
"narrower set of regex-anchored causal patterns (e.g., X causes Y, X leads "
"to Y, X requires Y) with hand-calibrated per-pattern confidence. Because "
"causal mentions in natural prose are frequently only two hops long and "
"fragmented across paraphrased surface forms, we further densify the causal "
"graph by (a) admitting high-confidence general-relationship triples whose "
"predicate is drawn from a curated causal/enablement verb set, and (b) merging "
"synonymous node surface forms via embedding-based canonicalization before "
"searching for longest paths - a lightweight, fully offline alternative to "
"learned event-causality models."
)
h2("2.4 Contradiction, Stance, and Conflict Detection")
p(
"Contradiction detection is closely related to natural language inference (NLI), "
"which typically classifies a hypothesis-premise pair as entailment, "
"contradiction, or neutral using a supervised classifier. Rather than deploying "
"a general-purpose NLI model - which can be brittle outside its training "
"distribution and offers no explanation for why two statements conflict - we "
"decompose contradiction detection into four inspectable, rule-governed "
"strategies operating over the already-extracted variables and relationships: "
"semantically-gated variable-value conflicts, negation/antonym relationship "
"conflicts, discourse-marker-triggered contrastive-claim detection, and "
"axis-based benefit-cost trade-off detection. The last of these is, to our "
"knowledge, an underexplored detection target: many technical documents assert "
"a benefit (precision, speed, efficiency) in one passage and acknowledge its "
"corresponding cost (off-target effects, safety risk, a limitation) in another, "
"without the two statements ever sharing a subject or object - a pattern the "
"first three strategies, by construction, cannot catch."
)
h2("2.5 Assumption and Premise Mining")
p(
"Argumentation mining and epistemic-stance detection literatures study how "
"authors signal certainty, hedging, and unstated premises via lexical and "
"syntactic cues (modal verbs, hedge words, conditional constructions). Our "
"assumption detector follows this lexical-cue tradition, combined with a "
"co-occurrence gate against a small set of domain-agnostic risk themes "
"(delivery, specificity, scalability, safety, accuracy, clinical/production use) "
"so that a bare modal verb does not, on its own, trigger a flagged assumption."
)
h2("2.6 Knowledge Graph Construction from Text")
p(
"Automatic knowledge-graph construction from unstructured text must address "
"entity resolution (merging surface-form variants of the same node), edge "
"sparsity (most extracted graphs are far sparser than the underlying semantic "
"graph), and - for interactive use - a principled way to keep a large extracted "
"graph within a renderable node budget. We address entity resolution with "
"agglomerative, cosine-similarity clustering over node-label embeddings; edge "
"sparsity with a transitive-closure and embedding-similarity-bridging inference "
"pass; and node-budget pruning with a salience-protection rule that exempts "
"problem participants, named entities, top-ranked concepts, and user-editable "
"domain seed terms from an otherwise pure degree-based cut. Community structure "
"is computed with greedy modularity maximization [8], as implemented in "
"NetworkX [9]."
)
h2("2.7 Retrieval-Augmented Question Answering and Hybrid Retrieval")
p(
"Retrieval-augmented generation typically pairs a dense or sparse retriever with "
"a generative reader. Our question-answering interface omits the generative "
"reader for its primary path - deterministic intent handlers answer common "
"query types (variable lookup, contact lookup, acronym definition, yes/no "
"problem queries, whole-paragraph retrieval, document summarization) directly "
"from the structured extraction - and falls back to a general retrieval path "
"only for open-ended questions. That general path follows established hybrid "
"information-retrieval practice: BM25 [5] lexical scoring and dense sentence-"
"embedding cosine similarity are combined via Reciprocal Rank Fusion, and the "
"resulting candidate pool is reranked with a cross-encoder [6] trained on the "
"MS MARCO passage-ranking dataset [7], a combination known in the information-"
"retrieval literature to outperform either signal alone."
)

# ── 3. System Overview ────────────────────────────────────────────────────────
h1("System Architecture Overview")
p(
"DocAnalyzer is a Flask-based web application. A single analysis request "
"proceeds through five stages: (i) multi-format ingestion and text "
"normalization; (ii) one shared spaCy parse producing sentence segmentation, "
"named entities, and subject-verb-object triples, which every downstream stage "
"consumes rather than re-parsing the document; (iii) a set of independent "
"extraction modules (typed entities, concepts, variables, structured data) that "
"run against the same parsed representation; (iv) a set of reasoning modules "
"(dependencies, causal chains, problems, contradictions, assumptions, knowledge "
"graph, executive insights) that consume the outputs of stage (iii); and (v) "
"presentation - a JSON result object rendered as a tabbed, resizable dual-panel "
"web interface, an exportable Markdown/Excel report, and a companion "
"question-answering endpoint over the same result object."
)
table(
 ["Stage", "Representative operations", "Primary technique(s)"],
 [
  ["Ingestion", "PDF / DOCX / Markdown / plain-text to normalized UTF-8 text",
   "markitdown; pdfplumber; python-docx"],
  ["Shared parse", "Sentence segmentation; named entities; SVO triples",
   "spaCy (en_core_web_lg)"],
  ["Typed entities", "Zero-shot job title / product / event spans",
   "GLiNER (gliner_small-v2.1)"],
  ["Concepts", "Keyphrase candidates; importance ranking; tiering; clustering",
   "KeyBERT (all-MiniLM-L6-v2)"],
  ["Variables / structured data",
   "Name=value patterns; molecular quantities; emails, URLs, acronyms, citations",
   "Regular expressions"],
  ["Relationships", "SVO triples; causal-pattern regex; per-triple confidence",
   "Dependency parse + regex + heuristic scoring"],
  ["Dependencies / causal chains",
   "Directed graph construction; cycle detection; longest-path search",
   "NetworkX"],
  ["Problems / contradictions / assumptions",
   "Severity tiering; four-strategy conflict detection; marker + thematic mining",
   "Rule-based, embedding-gated heuristics"],
  ["Knowledge graph", "Node canonicalization; salience-protected pruning; "
   "transitive/similarity inference; community detection",
   "Sentence embeddings + NetworkX"],
  ["Executive insights", "Deterministic template roll-up incl. competing-"
   "interests risk alert", "Rule-based"],
  ["Question answering", "Deterministic intent handlers; hybrid BM25/semantic/"
   "cross-encoder general retrieval", "rank_bm25 + Sentence-BERT + cross-encoder"],
 ],
 caption="Pipeline stages, representative operations, and primary techniques."
)
p(
"All core stages in Table 1 execute locally with no network dependency. A "
"single optional stage - AI-assisted problem/acronym enrichment via the "
"Anthropic API - is gated behind an environment-variable API key and is "
"skipped (with the deterministic result returned unchanged) whenever the key is "
"absent or the call fails; it is not required to reproduce any result reported "
"in this manuscript (see Section 7.3, Reproducibility)."
)

# ── 4. Methods ─────────────────────────────────────────────────────────────
h1("Methods")

h2("4.1 Document Ingestion and Preprocessing")
p(
"Uploaded files are converted to normalized text via format-specific readers "
"(pdfplumber for PDF, python-docx for Word, and a general Markdown/plain-text "
"path), concatenated across multiple uploaded files, and passed through a "
"text-cleaning stage that collapses whitespace, repairs common PDF-extraction "
"artifacts (ftfy), and isolates Markdown headings so that short, all-Title-Case "
"heading lines are not mis-parsed as sentences by the downstream NER stage. The "
"combined text is capped at 400,000 characters per analysis request."
)

h2("4.2 Named Entity Extraction and Tiering")
p(
"A single spaCy (en_core_web_lg) pass over the document yields sentence "
"boundaries, named entities, and dependency parses in one traversal, which "
"every subsequent stage reuses. Entities are filtered to remove PDF-extraction "
"noise, citation-year false positives, and bare numeric/percentage/money spans "
"(already captured by the structured-data layer, Section 4.5). Surviving "
"entities are ranked by an importance score and partitioned into three tiers by "
"rank percentile: Primary (top 20%), Secondary (next 35%), and Supporting "
"(remainder)."
)
p(
"spaCy's label set is generic (e.g., ORG, PERSON, GPE) and frequently "
"mis-classifies domain-specific terms. We re-type each entity by embedding its "
"surface text plus surrounding context (KeyBERT's underlying all-MiniLM-L6-v2 "
"encoder) and computing cosine similarity against a small library of "
"hand-written descriptions for each candidate domain label (e.g., PROTEIN, "
"GENE_SYSTEM, PATHWAY, INSTITUTION). A small override table takes precedence "
"for terms the embedding step is known to consistently mis-classify due to "
"name ambiguity."
)

h2("4.3 Zero-Shot Typed Entity Extraction")
p(
"GLiNER (gliner_small-v2.1) is queried in a zero-shot span-classification mode "
"for three labels not covered by spaCy's inventory - job title, product, and "
"event - over 1,200-character chunks (GLiNER's preferred window size) at a "
"minimum confidence threshold of 0.5. Duplicate spans are deduplicated by "
"(normalized text, label), keeping the highest-confidence occurrence."
)

h2("4.4 Key Concept Extraction")
p(
"KeyBERT (all-MiniLM-L6-v2) proposes keyphrase candidates from the document "
"text; a surrounding reasoning layer filters low-frequency or stopword-only "
"candidates, merges near-duplicate phrasings, and computes an importance score "
"combining KeyBERT's relevance score with frequency and entity co-occurrence. "
"Concepts are ranked and partitioned into three tiers by rank percentile - Core "
"(top 30%), Supporting (next 40%), Derived (remainder) - and grouped into "
"theme clusters for the executive-insight roll-up (Section 4.13)."
)

h2("4.5 Structured Data and Variable Extraction")
p(
"A library of regular-expression patterns extracts three families of "
"quantitative and structured content that a general-purpose NER model does not "
"reliably capture: (a) generic name=value / name: value constructions "
"(e.g., \"sequence length (nt) = 20\"); (b) domain-agnostic scientific "
"measurement patterns for density and biomass-style values; and (c) inline "
"molecular-biology quantities that do not fit a name=value template but are "
"the primary numerical claims in genomics-style text (e.g., \"~2,000 genes\", "
"\"20-nt spacer\", \"90% inheritance\"). A parallel structured-data extractor "
"identifies emails, URLs, phone numbers, percentages, currency amounts, "
"citations/DOIs, and in-text acronym definitions. Each extracted item carries "
"a source-sentence context string (up to 400 characters) for provenance."
)

h2("4.6 Relationship Extraction and Confidence Scoring")
p(
"General relationships are extracted as subject-verb-object (SVO) triples from "
"the dependency parse produced in the shared spaCy pass. Each triple is scored "
"by a confidence heuristic that starts from a baseline of 55 and adjusts for "
"signals available at parse time: +8 if the subject is a multi-word noun "
"phrase, +8 if the object is a multi-word noun phrase, +10 if the verb is not "
"drawn from a small list of semantically empty verbs (is, has, gets, ...), -8 "
"if either argument is trivially short (three characters or fewer), -10 for "
"sentences longer than 45 tokens (where dependency parses are more error-"
"prone), and +4 for sentences of 25 tokens or fewer; the result is clamped to "
"the range [40, 90]. A separate, narrower causal-relationship extractor applies "
"a curated set of causal regular-expression patterns (causes, leads to, "
"results in, requires, prevents, enables, ...), each with a hand-calibrated "
"confidence in the same numeric range. Relationships whose source sentence "
"contains an explicit negation adjacent to the extracted predicate are flagged "
"with a 'negated' hallucination-risk marker rather than silently reported as "
"asserted fact."
)

h2("4.7 Dependency Graph Construction")
p(
"Relationships whose predicate matches a curated set of dependency-indicating "
"verbs (depends on, requires, needs, relies on, ...) are compiled into a "
"directed graph. Simple-cycle detection over this graph flags nodes "
"participating in a circular dependency - i.e., a node that (directly or "
"indirectly) depends on itself - a pattern that is reported to the user as an "
"explicit 'Circular?' column and treated as a signal of fragile or tangled "
"reasoning."
)

h2("4.8 Causal Chain Construction")
p(
"Causal relationships alone are typically sparse (on the order of a dozen "
"pairs in a several-thousand-word document) and, taken at face value, form "
"mostly disconnected two-node edges. We address this in two ways. First, the "
"causal edge set is enriched with high-confidence (>= 75) general relationships "
"whose predicate verb is drawn from a curated causal/enablement whitelist "
"(enable, require, lead, cause, produce, allow, drive, trigger, generate, "
"provide, expedite, ensure, improve, reduce, increase, enhance, affect, "
"impact, and their inflections), so that a strong SVO relationship such as "
"'X enables Y' can extend a chain even when it was not tagged by the narrower "
"causal-pattern extractor. Second, node labels are canonicalized via the same "
"embedding-based agglomerative merging used for the knowledge graph (Section "
"4.12) before the graph is built, so that near-synonymous surface forms (e.g., "
"'efficiency of viral vector' and 'efficiency of viral vector packaging') "
"collapse to one node and no longer break an otherwise-continuous chain. Cycles "
"are removed, and for each root node (in-degree zero) an iterative "
"depth-first search finds the longest simple path to a sink; the resulting "
"chains are deduplicated, sorted by (length, mean edge confidence) descending, "
"and capped at 60."
)

h2("4.9 Problem and Risk Identification")
p(
"Every general relationship triple is scored against curated security/safety "
"and generic problem-keyword vocabularies. A triple whose subject, predicate, "
"and object jointly intersect the security vocabulary is tagged Critical "
"severity; a match against the broader problem vocabulary (gated by a "
"context check to avoid single-keyword false positives) is tagged High; "
"participation in a dependency cycle (Section 4.7) is tagged Medium; and a "
"problem-vocabulary match confined to the surrounding sentence context alone is "
"tagged Low. For each flagged problem, up to three root-cause ancestors are "
"looked up in the causal graph, and up to four co-mentioned named entities in "
"the same source sentence are recorded as affected entities. A direct sentence "
"scan additionally flags conservation/ecological risk phrasing not otherwise "
"captured by the relationship-based path."
)

h2("4.10 Contradiction Detection")
p(
"Four complementary, independently-triggered strategies populate the "
"contradiction table:"
)
numbered([
 "Variable-value conflicts. When the same variable name is assigned two "
 "different values, the two surrounding sentence contexts are embedded and "
 "compared by cosine similarity; a similarity below 0.60 indicates the two "
 "values describe different subjects sharing a coincidental label (e.g., a "
 "20-nt guide-target segment versus an 80-nt full synthesized guide) and is "
 "suppressed, while near-identical contexts (similarity > 0.95) are additionally "
 "checked for subset language ('including', 'of which', 'conserved across') "
 "that would indicate one value is a strict subset of the other rather than a "
 "genuine conflict. Surviving conflicts are assigned High or Medium severity by "
 "similarity strength.",
 "Negation and antonym conflicts. Relationship triples sharing the same "
 "(subject, object) pair are compared pairwise: identical predicates with "
 "differing negation status, or predicates that are curated antonym pairs "
 "(increase/decrease, enable/prevent, support/contradict, ...), are reported as "
 "High-severity conflicts.",
 "Contrastive discourse markers. Sentences containing an explicit contrastive "
 "discourse marker (however, nevertheless, in contrast, although, whereas, "
 "on the other hand, ...) together with a negation cue are paired with their "
 "immediately preceding sentence as an author-flagged tension.",
 "Benefit-cost trade-off axes. Three axes - precision versus off-target, "
 "speed/efficiency versus safety, and benefit versus limitation - are each "
 "defined by a pair of regular expressions for the upside and downside "
 "vocabulary. A single sentence asserting both sides is reported as a "
 "High-severity, author-acknowledged trade-off; two separate sentences, one "
 "per side, are reported as a Medium-severity trade-off. This strategy targets "
 "cross-cutting tensions that do not share a subject or object and are "
 "therefore invisible to strategies (1)-(3).",
])

h2("4.11 Assumption Detection")
p(
"Sentences are matched against an extensible table of assumption markers "
"(it is assumed, we assume, presumably, given that, should be, provided that, "
"in principle, it is conceivable, ...), each carrying a hand-assigned "
"confidence that determines whether the sentence is reported as an Explicit or "
"Implicit assumption. A second, thematic pass flags sentences that combine a "
"modal/expectation cue (will, would, should, expect, likely, sufficient, ...) "
"with a domain-agnostic risk theme (delivery, specificity, off-target, "
"scalability, efficiency, safety, accuracy, clinical/production use, dosage), "
"surfacing premises such as 'delivery methods will scale to human use' that a "
"pure marker-matching approach would miss because no single marker phrase is "
"present."
)

h2("4.12 Knowledge Graph Construction")
p(
"Relationship subjects and objects become graph nodes; node labels are "
"canonicalized by an agglomerative clustering procedure that merges two labels "
"when either their sentence-embedding cosine similarity exceeds 0.85 or one "
"label's token set is a subset of the other's, with the most frequent surface "
"form in each cluster chosen as the canonical label. Nodes are annotated with "
"entity type (where available) and a weight proportional to mention count; "
"edges carry the source relationship label, type, source sentence, and a "
"boolean flag for whether the edge also appears in the problems table. Because "
"a purely degree-based pruning rule to a fixed node budget (150 nodes) would "
"discard low-degree but conceptually salient terms, a protected-node set is "
"computed first - covering problem-table participants, all named entities, "
"the top 30 ranked key concepts, and a small, user-editable set of seed domain "
"terms - and exempted from the degree cut. The pruned graph is then densified "
"by two offline inference passes: transitive closure over a curated set of "
"transitive-safe predicates (depends on, requires, contains, leads to, "
"results in, causes), and similarity-bridging, which connects otherwise "
"disconnected components via their highest cross-component embedding "
"similarity. Community structure is computed by greedy modularity maximization "
"[8] over the undirected projection, and degree and betweenness centrality are "
"reported per node."
)

h2("4.13 Executive Insight Summarization")
p(
"A deterministic template layer condenses the analysis into a short, "
"human-readable insight list: the most important concept and primary entity, "
"the most severe problem, the most influential and highest-betweenness graph "
"nodes, the dominant theme cluster, counts of contradictions and high-"
"confidence assumptions, the root cause of the longest causal chain, and "
"recommended actions (e.g., prioritize mitigation of high/critical problems, "
"reconcile contradictions before acting). This layer additionally scans all "
"extracted relationships for competing- or conflicting-financial-interest "
"disclosure language; when present and not negated (e.g., 'declares no "
"competing interests'), it is surfaced as an explicit 'Risk - Competing "
"Interests' insight rather than left buried in the relationship table, on the "
"premise that governance-relevant disclosures merit a dedicated, hard-to-miss "
"flag in an audit-oriented tool."
)

h2("4.14 Hybrid-Retrieval Question Answering")
p(
"A companion question-answering endpoint accepts a free-text question about "
"the analyzed document. Before falling back to general retrieval, the "
"question is tried against an ordered battery of deterministic fast-path "
"intent handlers keyed on trigger phrases and structured-data lookups: "
"variable/count queries, contact (email/URL) queries, yes/no problem queries, "
"acronym definitions, entity lookups, contradiction queries, whole-paragraph "
"retrieval (Section 4.15), and document summarization. Only if no fast-path "
"handler fires does the pipeline construct a flattened knowledge base of "
"chunks spanning every analytical layer (raw sentences, entities, concepts, "
"problems, relationships, causal chains, insights, variables, structured "
"data), embed the question and every chunk, and rank chunks by a Reciprocal "
"Rank Fusion (k=60) of semantic cosine similarity and BM25 lexical score. The "
"top hybrid candidates are reranked by a cross-encoder (ms-marco-MiniLM-L-6-v2) "
"for a final relevance ordering, and the top match is presented as the primary "
"answer with supporting evidence grouped by analytical category."
)

h2("4.15 Full-Paragraph Retrieval")
p(
"Requests that explicitly ask for verbatim source text (full source paragraph, "
"exact passage, original wording, word-for-word, ...) are routed to a "
"dedicated handler that operates on whole paragraphs rather than single "
"sentences, so that the returned text is never truncated to a fragment. "
"Request boilerplate and stopwords are stripped from the question to isolate "
"its search topic (e.g., 'give the full source paragraph that mentions the "
"gene counts' reduces to 'gene counts'), and the resulting topic is used to "
"rank whole paragraphs by the same hybrid BM25/semantic/cross-encoder pipeline "
"described in Section 4.14. Paragraphs whose colon-terminated lead-in "
"introduces a numbered or bulleted list are automatically expanded to include "
"the list items that complete them, and the top two ranked paragraphs are "
"returned verbatim."
)

# ── 5. Implementation ────────────────────────────────────────────────────────
h1("Implementation")
p(
"The system is implemented in Python 3 as a single Flask application. Table 2 "
"lists the principal third-party components and the role each plays; all are "
"open-source and run on commodity CPU hardware without a GPU requirement, "
"though inference is faster with one available."
)
table(
 ["Component", "Role"],
 [
  ["spaCy (en_core_web_lg)", "Sentence segmentation, dependency parsing, base "
   "named entity recognition"],
  ["GLiNER (gliner_small-v2.1)", "Zero-shot typed entity extraction"],
  ["KeyBERT (all-MiniLM-L6-v2)", "Keyphrase/concept candidate generation and "
   "embedding backbone reused for entity re-typing, node canonicalization, and "
   "semantic retrieval"],
  ["sentence-transformers cross-encoder (ms-marco-MiniLM-L-6-v2)",
   "Reranking for question answering and paragraph retrieval"],
  ["rank_bm25", "Lexical (BM25) scoring for hybrid retrieval"],
  ["NetworkX", "Directed/undirected graph construction, cycle detection, "
   "longest-path search, community detection, centrality"],
  ["scikit-learn", "Cosine similarity computation"],
  ["markitdown, pdfplumber, python-docx", "Multi-format document ingestion"],
  ["pandas, openpyxl", "Markdown/Excel report export"],
  ["Anthropic API (optional)", "Supplementary, disclosed problem/acronym "
   "enrichment; not required for core functionality"],
 ],
 caption="Principal software components and their role in the pipeline."
)
p(
"The web front end is a single-page vanilla-JavaScript interface. Results are "
"presented in two coordinated regions: a horizontal tab bar for document-level "
"risk categories (Problems, Dependencies, Contradictions, Assumptions), and a "
"resizable dual-panel split view separating higher-level analysis (Key "
"Concepts, Relationships, Causal Chains, Insights, Knowledge Graph) from raw "
"extracted layers (Entities, Typed Entities, Structured Data, Variables). "
"Column order within each table is fixed per tab (rather than inferred from "
"row key order) so that the most decision-relevant field is always presented "
"first, and column headers for non-obvious metrics (e.g., 'Circular?', "
"'Hallucination Risk', 'Avg Confidence') carry an inline hover explanation. "
"The knowledge graph is rendered as an interactive force-directed network "
"(vis.js) with node size proportional to connectivity and color keyed to "
"detected community."
)

# ── 6. Case Study ─────────────────────────────────────────────────────────────
h1("Case Study: Qualitative Evaluation")
h2("6.1 Corpus")
p(
"We applied the full pipeline to a two-document corpus totaling 11,047 words: "
"an introductory background text (crispr_microbiology_background.md, 642 "
"words) and a peer-reviewed review article, Barrangou and Doudna, "
"'Applications of CRISPR technologies in research and beyond,' Nature "
"Biotechnology 34(9):933-941, 2016 [10] (10,405 words). This corpus was chosen "
"because it combines an accessible tutorial register with a dense, "
"citation-heavy scientific-review register in the same analysis run, and "
"because the review article's own text supplied several of the ground-truth "
"reference points used to develop and debug the extraction rules described in "
"Section 4 over the course of this project."
)
h2("6.2 Summary Statistics")
p(
"Table 3 reports the summary metrics produced by an end-to-end run of the "
"pipeline on the corpus described above."
)
table(
 ["Metric", "Value"],
 [
  ["Files analyzed", "2"],
  ["Total words", "11,047"],
  ["Named entities", "59"],
  ["Typed entities", "16"],
  ["Structured data items", "23"],
  ["Key concepts", "13"],
  ["Variables detected", "5"],
  ["Relationships mapped", "179"],
  ["Dependencies", "22"],
  ["Causal chains", "30"],
  ["Problems", "7"],
  ["Contradictions", "8"],
  ["Assumptions", "15"],
  ["Knowledge graph nodes", "149"],
  ["Knowledge graph edges", "128"],
 ],
 caption="Summary statistics for the two-document case-study corpus (single "
         "end-to-end run reflecting all Section 4 enhancements)."
)
p(
"This run reflects a single end-to-end execution of the complete pipeline "
"described in Section 4, including the causal-chain densification, "
"salience-protected graph pruning, benefit-cost trade-off contradiction axis, "
"thematic-assumption pass, and competing-interests executive-insight alert "
"(Sections 4.8, 4.10, 4.11, 4.13). An earlier run of the same corpus, captured "
"before those four enhancements were implemented, is retained here as a "
"before/after comparison: causal chains increased from 15 (maximum length 2) "
"to 30, of which 5 now reach length 3 via the causal-verb-enriched, "
"canonicalized edge set; contradictions increased from 5 to 8 with the "
"addition of three benefit-cost trade-off detections; assumptions increased "
"from 2 to 15, of which 10 were surfaced only by the new thematic "
"co-occurrence pass (Section 4.11) and would not have been flagged by marker "
"matching alone; and the knowledge graph retained 149 nodes and 128 edges "
"against a raised 150-node budget, up from 119 nodes and 102 edges under the "
"prior 120-node budget. The competing-interests detector (Section 4.13) fired "
"on this corpus as intended, surfacing the source article's financial "
"disclosure statement as a dedicated executive insight rather than leaving it "
"embedded in the relationships table (Section 6.3)."
)
h2("6.3 Representative Extraction Examples")
p(
"Table 4 gives one representative example per analytical layer, quoted "
"directly from the case-study run, to illustrate output granularity and "
"provenance."
)
table(
 ["Layer", "Example output"],
 [
  ["Named entity", "CRISPR - GENE_SYSTEM, 57 occurrences, tier Primary"],
  ["Typed entity", "\"CRISPR-Cas9\" - product, confidence 0.847"],
  ["Structured data", "Acronym: PAM = \"protospacer adjacent motif\""],
  ["Variable", "gene count = \"2,000 genes\" (molecular quantity)"],
  ["Relationship", "\"Guide RNA\" -directs-> \"Cas9\" (syntactic SVO, "
   "confidence 77)"],
  ["Dependency", "\"cas proteins\" depends on \"guide rna\""],
  ["Causal chain", "\"alternative end joining/microhomology-mediated "
   "joining/mismatch and nucleotide-excision repair\" -> \"CRISPR-mediated "
   "genome editing\" -> \"development of large animal models of...\" "
   "(length 3, confidence 83.0)"],
  ["Problem", "[High] \"Current limitations include: - Off-target mutations - "
   "Delivery difficulties - Ethical concerns ...\""],
  ["Contradiction (contrastive)", "Medium: ZFN/TALEN protein-engineering "
   "burden versus CRISPR's guide-synthesis-only requirement"],
  ["Contradiction (trade-off)", "High, precision vs off-target: \"It allows "
   "scientists to make precise modifications to DNA sequences...\" versus "
   "\"...PAM-constrained targeting and off-target activity can affect "
   "editing precision\""],
  ["Assumption (thematic)", "Implicit (\"expectation about safety\", "
   "confidence 55): continued research will likely expand the capabilities "
   "and safety of CRISPR-based technologies"],
  ["Executive insight (competing interests)", "Risk - Competing Interests: "
   "financial conflict-of-interest disclosed - verify independence: "
   "\"The authors declare competing financial interests...\""],
  ["Executive insight (root cause)", "Root Cause: the alternative-end-joining "
   "/ microhomology-mediated-joining repair-pathway statement initiates the "
   "longest causal chain"],
 ],
 caption="Representative extraction output, one example per analytical layer "
         "(single end-to-end run reflecting all Section 4 enhancements)."
)
h2("6.4 Question-Answering Behavior")
p(
"Four illustrative questions were posed against the analyzed corpus. A "
"grounding question ('Does the document mention any clinical trials for "
"CRISPR-based cancer therapies in Europe?') correctly retrieved the passage "
"stating that two such trials were approved in China and the United States "
"and did not fabricate a European trial the source text does not support - a "
"desirable property for an audit-oriented tool, where declining to assert an "
"unsupported claim is preferable to a fluent but ungrounded answer. A "
"paragraph-retrieval question ('Give the full source paragraph that mentions "
"the gene counts') correctly returned the paragraph describing the "
"CRISPR-based screen that identified approximately 2,000 core fitness genes, "
"verbatim and in full, rather than a single truncated sentence."
)
p(
"A broader, more open-ended question ('What are the main non-medical "
"applications of CRISPR mentioned in the documents?') was intercepted by the "
"deterministic entity-lookup fast-path handler (Section 4.14) and answered "
"with a generic entity description rather than a synthesized list of the "
"agricultural, food-science, and antimicrobial applications the source text "
"actually discusses at length. This is a genuine limitation of the current "
"intent-routing order rather than a retrieval failure: the fast-path handlers "
"are intentionally tried before general retrieval to guarantee fast, exact "
"answers for the query types they cover, but a broad, multi-topic question can "
"still match a narrower handler's trigger pattern before it reaches the "
"hybrid-retrieval path that would otherwise assemble a more complete answer. "
"We discuss this trade-off further in Section 7.1."
)

# ── 7. Discussion ─────────────────────────────────────────────────────────────
h1("Discussion")

h2("7.1 Design Rationale and Observed Trade-offs")
p(
"The central design decision in this system is to prefer many small, "
"inspectable, rule-governed extraction and detection modules over a single "
"end-to-end learned or LLM-driven model. This yields determinism (the same "
"document produces the same output on every run), auditability (every "
"extracted item traces to a specific rule, pattern, or confidence "
"computation), and independence from a paid or rate-limited external API for "
"the core pipeline. The cost, illustrated concretely in Section 6.4, is that "
"an ordered battery of narrow intent handlers can occasionally intercept a "
"broad question before a more holistic retrieval path would have produced a "
"better answer; and that rule-governed detectors (Section 4.10-4.11) trade "
"recall for precision relative to a general-purpose NLI or LLM-based "
"classifier, by design."
)

h2("7.2 Limitations and Threats to Validity")
bullets([
 "No labeled benchmark. The case study in Section 6 is a single, "
 "qualitatively-inspected corpus. We have not constructed a labeled reference "
 "set against which precision, recall, or F1 can be computed for any "
 "extraction layer, and no claim in this manuscript should be read as a "
 "quantitative accuracy claim.",
 "Single corpus, single run. Section 6.2 reports one before/after pair of "
 "end-to-end runs on one corpus. This demonstrates that the Section 4 "
 "enhancements function together without regressing the pipeline, but it is "
 "not a substitute for the repeated-run and multi-corpus reproducibility "
 "checks called for elsewhere in this section.",
 "Single-domain corpus. The case study is drawn from biomedical/genome-editing "
 "text. Several extraction components (the molecular-quantity regex library, "
 "the risk-theme vocabulary used by the thematic assumption pass, the "
 "trade-off axis vocabularies) were developed and tuned against this domain "
 "and its vocabulary; generalization to other domains (legal, financial, "
 "engineering) has not been evaluated and may require additional or "
 "re-weighted lexical resources.",
 "Confidence scores are heuristic, not calibrated. The per-triple relationship "
 "confidence (Section 4.6) and per-pattern causal confidence are hand-"
 "calibrated numeric scores intended to rank extractions relative to one "
 "another; they are not probabilities calibrated against ground truth and "
 "should not be interpreted as such.",
 "Fast-path/retrieval interaction. As illustrated in Section 6.4, the "
 "ordered intent-handler design can occasionally produce a narrower answer "
 "than the general retrieval path would have for broad, multi-topic "
 "questions.",
 "Embedding-based canonicalization can over- or under-merge. Cosine-"
 "similarity node merging (Sections 4.8, 4.12) is threshold-based and can, in "
 "principle, either merge two genuinely distinct entities with similar "
 "phrasing or fail to merge two paraphrases of the same entity; we have not "
 "measured the rate of either error.",
])

h2("7.3 Reproducibility")
p(
"Every result reported in Section 6 was produced by the deterministic core "
"pipeline described in Sections 4.1-4.15, none of which requires network "
"access or an API key. The single optional cloud-LLM enrichment stage "
"(Section 3) is disclosed, off by default in the absence of an API key, and "
"was not used to produce any figure in this manuscript. Re-running the "
"pipeline on the same input document with the same library versions should "
"reproduce the same entities, relationships, and structured data; the "
"knowledge-graph community-detection step (greedy modularity maximization) "
"and any downstream tie-breaking are, to the best of our knowledge, "
"deterministic given a fixed NetworkX version, but this has not been "
"independently verified across library versions and is noted here as an item "
"for the reproducibility checklist accompanying formal submission."
)

h2("7.4 Ethical Considerations")
p(
"A tool that automatically flags 'problems,' 'contradictions,' and "
"'assumptions' in a source document carries a risk of over-claiming: a "
"severity label of Critical or High is a heuristic classification, not an "
"authoritative judgment about the source document's scientific validity, and "
"should be presented to end users with that caveat. The competing-interests "
"insight (Section 4.13) is a governance-relevant example of this same "
"concern in the opposite direction: surfacing a disclosure prominently is "
"intended to aid a human reviewer, not to imply misconduct, and the detector "
"is explicitly negation-aware so that a normal 'no competing interests' "
"disclosure is not mis-flagged as a risk."
)

# ── 8. Conclusion and Future Work ─────────────────────────────────────────────
h1("Conclusion and Future Work")
p(
"We have described DocAnalyzer, a fully offline, modular pipeline that "
"extracts eleven linked analytical layers from an arbitrary input document, "
"and a companion hybrid-retrieval question-answering interface, without "
"requiring a paid or network-dependent large language model for any "
"core-pipeline result. A qualitative case study on a peer-reviewed "
"genome-editing review article demonstrates the system's output across every "
"layer and surfaces concrete, honestly-reported limitations - most notably "
"the current absence of a labeled benchmark corpus."
)
p("Immediate priorities before this manuscript is submitted for review are:")
numbered([
 "Construction of a small labeled evaluation set (entity, relationship, and "
 "contradiction annotations by at least two independent annotators with "
 "inter-annotator agreement reported) sufficient to report precision/recall "
 "for the highest-stakes extraction layers (problems, contradictions, "
 "assumptions).",
 "Evaluation on at least one additional, non-biomedical document domain to "
 "assess generalization of the domain-tunable vocabularies identified in "
 "Section 7.2.",
 "A user study measuring whether the fast-path/retrieval intent-routing "
 "design (Section 7.1) helps or hinders real users on open-ended questions, "
 "relative to always routing through general hybrid retrieval.",
 "Repetition of the Section 6 before/after comparison on additional corpora "
 "to confirm that the causal-chain, contradiction, assumption, and "
 "graph-retention gains observed in this single run generalize.",
])

# ── Acknowledgments ───────────────────────────────────────────────────────────
pagebreak()
h1("Acknowledgments", numbered=False)
p("[To be completed by the author(s).]", italic=True)

# ── References ────────────────────────────────────────────────────────────────
h1("References", numbered=False)
p(
"Citations below correspond to the underlying methods, models, and datasets "
"employed by the pipeline components described in Sections 4-5, and to the "
"case-study source document (Section 6). Exact citation formatting, volume/"
"page details, and completeness should be verified and adjusted to the "
"target journal's required citation style before submission.",
 italic=True
)
refs = [
"[1] Honnibal, M., & Montani, I. (2017). spaCy 2: Natural language "
"understanding with Bloom embeddings, convolutional neural networks and "
"incremental parsing.",
"[2] Grootendorst, M. (2020). KeyBERT: Minimal keyword extraction with BERT. "
"Zenodo. https://doi.org/10.5281/zenodo.4461265",
"[3] Zaratiana, U., Tomeh, N., Holat, P., & Charnois, T. (2023). GLiNER: "
"Generalist model for named entity recognition using bidirectional "
"transformer. arXiv:2311.08526.",
"[4] Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings "
"using Siamese BERT-networks. Proceedings of EMNLP-IJCNLP 2019, 3982-3992.",
"[5] Robertson, S., & Zaragoza, H. (2009). The probabilistic relevance "
"framework: BM25 and beyond. Foundations and Trends in Information "
"Retrieval, 3(4), 333-389.",
"[6] Nogueira, R., & Cho, K. (2019). Passage re-ranking with BERT. "
"arXiv:1901.04085.",
"[7] Bajaj, P., Campos, D., Craswell, N., et al. (2016). MS MARCO: A human "
"generated machine reading comprehension dataset. arXiv:1611.09268.",
"[8] Clauset, A., Newman, M. E. J., & Moore, C. (2004). Finding community "
"structure in very large networks. Physical Review E, 70(6), 066111.",
"[9] Hagberg, A. A., Schult, D. A., & Swart, P. J. (2008). Exploring network "
"structure, dynamics, and function using NetworkX. Proceedings of the 7th "
"Python in Science Conference (SciPy2008), 11-15.",
"[10] Barrangou, R., & Doudna, J. A. (2016). Applications of CRISPR "
"technologies in research and beyond. Nature Biotechnology, 34(9), 933-941. "
"https://doi.org/10.1038/nbt.3659",
]
for r in refs:
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.first_line_indent = Inches(-0.3)
    para.add_run(r)

# ── Appendices ────────────────────────────────────────────────────────────────
pagebreak()
h1("Appendix A: Result Schema Reference", numbered=False)
p(
"Every analysis produces a single JSON result object; Table 5 lists its "
"top-level categories and the fields reported for each row."
)
table(
 ["Category", "Representative fields"],
 [
  ["entities", "Entity, Type, Description, Tier, Importance, Occurrences, "
   "Context"],
  ["typed_entities", "Entity, Label, Confidence, Context"],
  ["concepts", "Concept, Type, Importance, Relevance Score, Theme/Cluster, "
   "Frequency, Related Entities, Description"],
  ["structured", "Type, Value, Detail, Context"],
  ["variables", "Variable/Parameter, Value, Type, Context"],
  ["relationships", "Subject, Relationship, Object, Rel. Type, Confidence, "
   "Source Sentence, Hallucination Risk"],
  ["problems", "Problem Type, Object, Affected Entities, Subject, "
   "Relationship, Severity, Root Causes, Source Sentence"],
  ["dependencies", "Node, Depends On, Direct Deps, Total (incl. indirect), "
   "Strength, Circular?"],
  ["causal_chains", "Causal Chain, Length, Root Cause, Final Outcome, Avg "
   "Confidence"],
  ["contradictions", "Statement A, Statement B, Conflict Type, Severity"],
  ["assumptions", "Assumption, Type, Marker, Confidence, Impact if False"],
  ["insights", "Category, Finding"],
  ["graph", "nodes[] (id, label, type, size, centrality, betweenness, "
   "influence, group), edges[] (from, to, label, rtype, sentence, problem, "
   "inferred, weight), stats"],
  ["summary", "Per-category counts and model identifiers"],
 ],
 caption="Top-level result-object categories and representative fields."
)

h1("Appendix B: Reproducibility Environment", numbered=False)
p(
"The pipeline was implemented and the case study in Section 6 was produced "
"using the Python package versions (minimum bounds) listed below. Exact "
"pinned versions used for the final submitted case-study run should be "
"recorded via 'pip freeze' at the time that run is executed."
)
env_pkgs = [
 "flask>=3.0.0", "spacy>=3.7.0 (model: en_core_web_lg)", "keybert>=0.8.0",
 "sentence-transformers>=2.2.0", "markitdown[all]>=0.0.2", "gliner>=0.2.0 "
 "(model: urchade/gliner_small-v2.1)", "phonenumbers>=9.0.0",
 "tldextract>=5.0.0", "quantulum3>=0.10.0", "dateparser>=1.2.0",
 "price-parser>=0.3.0", "validators>=0.30.0", "scikit-learn>=1.3.0",
 "pdfplumber>=0.10.0", "python-docx>=1.1.0", "pandas>=2.0.0",
 "openpyxl>=3.1.0", "networkx>=3.0", "nltk>=3.8.0", "numpy>=1.24.0",
 "ftfy>=6.0.0", "anthropic>=0.40.0 (optional, disclosed enrichment step "
 "only)", "rank-bm25>=0.2.2",
]
bullets(env_pkgs)

doc.save("Technical_Whitepaper.docx")
print("Wrote Technical_Whitepaper.docx")
